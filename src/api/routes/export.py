"""
Export Routes - Document and data export endpoints
"""
import logging
import os
import re
import tempfile
from datetime import datetime
from enum import Enum
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

from src.agents.period_insight_agent import PeriodInsightAgent
from src.api.dependencies import load_dashboard_data
from src.tools.chart_generator import ChartGenerator
from src.tools.external_signal_collector import ExternalSignalCollector
from src.tools.period_analyzer import PeriodAnalyzer
from src.tools.reference_tracker import ReferenceTracker
from src.tools.sqlite_storage import get_sqlite_storage

router = APIRouter(prefix="/api/export", tags=["export"])


class SignalRelevance(Enum):
    """외부 신호 시간적 관련성 분류"""
    TIER1_CORE = "core"           # 분석 기간 ±7일 - 직접 관련
    TIER2_BACKGROUND = "background"  # ±30일 또는 구조적 트렌드
    TIER3_ARCHIVE = "archive"     # 30일+ 이전, 장기 트렌드만


# 구조적 트렌드 키워드 (시간 민감도 낮음)
STRUCTURAL_TREND_KEYWORDS = [
    "trend", "growth", "market size", "industry", "expansion",
    "k-beauty", "clean beauty", "sustainable", "global",
    "market report", "forecast", "analysis"
]


class ExportRequest(BaseModel):
    """내보내기 요청"""
    start_date: str | None = None
    end_date: str | None = None
    include_strategy: bool = True
    include_external_signals: bool = True  # External Signal 포함 여부


class AnalystReportRequest(BaseModel):
    """애널리스트 리포트 요청"""
    start_date: str  # Required: YYYY-MM-DD
    end_date: str    # Required: YYYY-MM-DD
    include_charts: bool = True
    include_external_signals: bool = True


def _classify_signal_relevance(
    signal,
    analysis_start: datetime,
    analysis_end: datetime
) -> SignalRelevance:
    """
    외부 신호의 시간적 관련성 분류 (3-Tier)

    Args:
        signal: ExternalSignal 객체
        analysis_start: 분석 시작일
        analysis_end: 분석 종료일

    Returns:
        SignalRelevance enum
    """
    # 신호 날짜 파싱
    signal_date = None
    published_at = getattr(signal, 'published_at', None)

    if published_at:
        try:
            if isinstance(published_at, str):
                # ISO 형식 또는 다양한 형식 처리
                published_at = published_at.replace('Z', '+00:00')
                if 'T' in published_at:
                    signal_date = datetime.fromisoformat(published_at).replace(tzinfo=None)
                else:
                    signal_date = datetime.strptime(published_at[:10], "%Y-%m-%d")
            elif isinstance(published_at, datetime):
                signal_date = published_at.replace(tzinfo=None)
        except Exception:
            signal_date = None

    # 구조적 트렌드 여부 확인
    title = getattr(signal, 'title', '').lower()
    content = getattr(signal, 'content', '').lower()
    combined_text = title + " " + content

    is_structural = any(
        keyword.lower() in combined_text
        for keyword in STRUCTURAL_TREND_KEYWORDS
    )

    # 날짜 기반 분류
    if signal_date:
        days_from_end = (signal_date.date() - analysis_end.date()).days

        # TIER 1: 분석 기간 ±7일
        if -7 <= days_from_end <= 7:
            return SignalRelevance.TIER1_CORE

        # TIER 2: ±30일 또는 구조적 트렌드
        if -30 <= days_from_end <= 30 or is_structural:
            return SignalRelevance.TIER2_BACKGROUND

        # TIER 3: 30일+ 이전이지만 구조적 트렌드인 경우만
        if is_structural:
            return SignalRelevance.TIER3_ARCHIVE

        # 이벤트성 뉴스 30일+ 외는 None 반환 (제외 대상)
        return None

    # 날짜 없으면 구조적 트렌드인 경우만 TIER 2
    if is_structural:
        return SignalRelevance.TIER2_BACKGROUND

    # 날짜 없고 구조적도 아니면 TIER 2로 분류 (보수적 접근)
    return SignalRelevance.TIER2_BACKGROUND


async def _get_external_signals(
    days: int = 7,
    brands: list = None,
    include_tavily: bool = True,
    start_date: str = None,
    end_date: str = None
) -> dict:
    """
    External Signal 수집 및 3-Tier 분류

    Args:
        days: 검색 기간 (일)
        brands: 검색할 브랜드 리스트
        include_tavily: Tavily 뉴스 검색 포함 여부
        start_date: 분석 시작일 (YYYY-MM-DD) - 3-Tier 분류용
        end_date: 분석 종료일 (YYYY-MM-DD) - 3-Tier 분류용

    Returns:
        {
            "signals": List[ExternalSignal],
            "classified": {
                "tier1_core": List[ExternalSignal],
                "tier2_background": List[ExternalSignal],
                "tier3_archive": List[ExternalSignal]
            },
            "report_section": str
        }
    """
    logger = logging.getLogger(__name__)
    all_signals = []

    try:
        collector = ExternalSignalCollector()
        await collector.initialize()

        # 브랜드 기본값
        if not brands:
            brands = ["LANEIGE", "COSRX", "K-Beauty"]

        # 1. Tavily 뉴스 검색 (최적화)
        if include_tavily:
            try:
                # 검색 기간: 최소 14일, 최대 30일 (분석 기간에 맞춤)
                search_days = min(max(days, 14), 30)
                tavily_signals = await collector.fetch_tavily_news(
                    brands=brands[:3],
                    topics=["K-Beauty skincare", "Amazon beauty bestseller"],
                    days=search_days,
                    max_results=15  # 더 많은 결과 수집
                )
                all_signals.extend(tavily_signals)
                logger.info(f"Collected {len(tavily_signals)} Tavily news signals (search_days={search_days})")
            except Exception as e:
                logger.warning(f"Tavily news fetch failed: {e}")

        # 2. RSS 피드 수집 (기존)
        try:
            rss_signals = await collector.fetch_all_rss_feeds(
                keywords=brands + ["skincare", "lip care"]
            )
            all_signals.extend(rss_signals)
            logger.info(f"Collected {len(rss_signals)} RSS signals")
        except Exception as e:
            logger.warning(f"RSS fetch failed: {e}")

        # 3. Reddit 트렌드 수집 (기존)
        try:
            reddit_signals = await collector.fetch_reddit_trends()
            all_signals.extend(reddit_signals)
            logger.info(f"Collected {len(reddit_signals)} Reddit signals")
        except Exception as e:
            logger.warning(f"Reddit fetch failed: {e}")

        # 기존 수집된 신호도 추가
        if collector.signals:
            all_signals.extend(collector.signals)

        # 중복 제거 (URL 기준)
        seen_urls = set()
        unique_signals = []
        for signal in all_signals:
            url = getattr(signal, 'url', '')
            if url and url not in seen_urls:
                seen_urls.add(url)
                unique_signals.append(signal)
            elif not url:
                unique_signals.append(signal)

        # 3-Tier 분류
        classified = {
            "tier1_core": [],      # 본문 인용 + 참고자료
            "tier2_background": [], # 참고자료만
            "tier3_archive": []    # 배경 자료 섹션
        }

        if start_date and end_date:
            try:
                analysis_start = datetime.strptime(start_date, "%Y-%m-%d")
                analysis_end = datetime.strptime(end_date, "%Y-%m-%d")

                filtered_signals = []
                for signal in unique_signals:
                    relevance = _classify_signal_relevance(signal, analysis_start, analysis_end)

                    if relevance == SignalRelevance.TIER1_CORE:
                        classified["tier1_core"].append(signal)
                        filtered_signals.append(signal)
                    elif relevance == SignalRelevance.TIER2_BACKGROUND:
                        classified["tier2_background"].append(signal)
                        filtered_signals.append(signal)
                    elif relevance == SignalRelevance.TIER3_ARCHIVE:
                        classified["tier3_archive"].append(signal)
                        # TIER3는 참고자료에만 포함, 본문 분석에서는 제외
                    # relevance가 None이면 제외 (이벤트성 + 30일+ 외)

                logger.info(
                    f"Signal classification: TIER1={len(classified['tier1_core'])}, "
                    f"TIER2={len(classified['tier2_background'])}, "
                    f"TIER3={len(classified['tier3_archive'])}"
                )

                # 필터링된 신호로 교체 (TIER3 제외)
                unique_signals = filtered_signals

            except Exception as e:
                logger.warning(f"Signal classification failed: {e}")
                # 분류 실패 시 전체를 TIER2로 처리
                classified["tier2_background"] = unique_signals

        # 보고서 섹션 생성
        if unique_signals:
            # 컬렉터에 신호 추가 후 보고서 생성
            for signal in unique_signals:
                if signal not in collector.signals:
                    collector.signals.append(signal)
            report_section = collector.generate_report_section(days=days)
        else:
            report_section = ""

        return {
            "signals": unique_signals,
            "classified": classified,
            "report_section": report_section
        }

    except Exception as e:
        logger.warning(f"External signal collection failed: {e}")
        return {"signals": [], "classified": {"tier1_core": [], "tier2_background": [], "tier3_archive": []}, "report_section": ""}
    finally:
        try:
            await collector.close()
        except Exception:
            pass


@router.post("/docx")
async def export_docx(request: ExportRequest):
    """
    인사이트 리포트 DOCX 생성 및 다운로드
    """
    data = load_dashboard_data()
    if not data:
        raise HTTPException(status_code=404, detail="Dashboard data not found")

    # DOCX 문서 생성
    doc = Document()

    # 스타일 설정
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # ===== 표지 =====
    title = doc.add_heading('AMORE INSIGHT Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('LANEIGE Amazon US 분석 리포트')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 날짜
    metadata = data.get("metadata", {})
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(f"분석 기준일: {metadata.get('data_date', datetime.now().strftime('%Y-%m-%d'))}")
    date_para.add_run(f"\n생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_page_break()

    # ===== 목차 =====
    doc.add_heading('목차', 1)
    toc_items = [
        "1. 요약 통계",
        "2. 브랜드별 성과",
        "3. 카테고리별 분석",
        "4. 주요 제품",
        "5. AI 인사이트 및 전략 제언",
        "6. 외부 트렌드 신호"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_page_break()

    # ===== 1. 요약 통계 =====
    doc.add_heading('1. 요약 통계', 1)
    summary = data.get("summary", {})

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    stats = [
        ("총 제품 수", summary.get("total_products", 0)),
        ("크롤링 카테고리", summary.get("categories_count", 0)),
        ("LANEIGE 제품 수", summary.get("laneige_products", 0)),
        ("평균 가격", f"${summary.get('avg_price', 0):.2f}"),
        ("데이터 날짜", metadata.get("data_date", "N/A"))
    ]

    for i, (label, value) in enumerate(stats):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # ===== 2. 브랜드별 성과 =====
    doc.add_heading('2. 브랜드별 성과', 1)

    brand_data = data.get("brand_metrics", [])
    if brand_data:
        # 상위 10개 브랜드
        top_brands = sorted(brand_data, key=lambda x: x.get("product_count", 0), reverse=True)[:10]

        brand_table = doc.add_table(rows=len(top_brands) + 1, cols=4)
        brand_table.style = 'Light Grid Accent 1'

        # 헤더
        headers = brand_table.rows[0].cells
        headers[0].text = "브랜드"
        headers[1].text = "제품 수"
        headers[2].text = "평균 순위"
        headers[3].text = "SoS (%)"

        # 데이터
        for i, brand in enumerate(top_brands, start=1):
            cells = brand_table.rows[i].cells
            cells[0].text = brand.get("brand", "Unknown")
            cells[1].text = str(brand.get("product_count", 0))
            cells[2].text = f"{brand.get('avg_rank', 0):.1f}"
            cells[3].text = f"{brand.get('sos', 0):.2f}%"
    else:
        doc.add_paragraph("브랜드 데이터가 없습니다.")

    doc.add_page_break()

    # ===== 3. 카테고리별 분석 =====
    doc.add_heading('3. 카테고리별 분석', 1)

    category_data = data.get("category_metrics", [])
    if category_data:
        for category in category_data:
            cat_name = category.get("category", "Unknown")
            doc.add_heading(cat_name, 2)

            cat_stats = [
                ("총 제품 수", category.get("total_products", 0)),
                ("LANEIGE 제품 수", category.get("laneige_products", 0)),
                ("평균 가격", f"${category.get('avg_price', 0):.2f}"),
                ("HHI", f"{category.get('hhi', 0):.2f}"),
                ("CPI", f"{category.get('cpi', 0):.2f}")
            ]

            for label, value in cat_stats:
                doc.add_paragraph(f"{label}: {value}")

            doc.add_paragraph()
    else:
        doc.add_paragraph("카테고리 데이터가 없습니다.")

    doc.add_page_break()

    # ===== 4. 주요 제품 =====
    doc.add_heading('4. 주요 제품 (LANEIGE Top 10)', 1)

    products = data.get("products", [])
    laneige_products = [p for p in products if p.get("brand", "").upper() == "LANEIGE"]
    laneige_products = sorted(laneige_products, key=lambda x: x.get("rank", 999))[:10]

    if laneige_products:
        product_table = doc.add_table(rows=len(laneige_products) + 1, cols=5)
        product_table.style = 'Light Grid Accent 1'

        # 헤더
        headers = product_table.rows[0].cells
        headers[0].text = "순위"
        headers[1].text = "제품명"
        headers[2].text = "카테고리"
        headers[3].text = "가격"
        headers[4].text = "평점"

        # 데이터
        for i, product in enumerate(laneige_products, start=1):
            cells = product_table.rows[i].cells
            cells[0].text = str(product.get("rank", "N/A"))
            cells[1].text = product.get("title", "Unknown")[:50]
            cells[2].text = product.get("category", "Unknown")
            cells[3].text = f"${product.get('price', 0):.2f}" if product.get("price") else "N/A"
            cells[4].text = str(product.get("rating", "N/A"))
    else:
        doc.add_paragraph("LANEIGE 제품이 없습니다.")

    doc.add_page_break()

    # ===== 5. AI 인사이트 및 전략 제언 =====
    if request.include_strategy:
        doc.add_heading('5. AI 인사이트 및 전략 제언', 1)

        insights = data.get("ai_insights", {})
        strategic_insights = insights.get("strategic_insights", [])

        if strategic_insights:
            for insight in strategic_insights:
                doc.add_heading(insight.get("title", "Insight"), 2)
                doc.add_paragraph(insight.get("content", ""))
                doc.add_paragraph()
        else:
            # 폴백 전략
            doc.add_paragraph("""
1. Top 10 유지 전략: 현재 상위권 제품의 리뷰 관리 및 재고 확보를 통한 포지션 유지

2. 경쟁사 모니터링: e.l.f., Maybelline 등 주요 경쟁사의 가격 및 프로모션 동향 파악

3. 신규 진입 기회: Lip Care 카테고리 외 Face Powder, Toner 등 확장 가능성 검토
""")

    # ===== 6. 외부 트렌드 신호 =====
    if request.include_external_signals:
        doc.add_heading('6. 외부 트렌드 신호', 1)

        # 분석 기간 계산
        if request.start_date and request.end_date:
            try:
                start = datetime.strptime(request.start_date, "%Y-%m-%d")
                end = datetime.strptime(request.end_date, "%Y-%m-%d")
                days = (end - start).days + 1
            except ValueError:
                days = 7
        else:
            days = 7

        signals_result = await _get_external_signals(
            days=days,
            start_date=request.start_date,
            end_date=request.end_date
        )
        external_signals_text = signals_result.get("report_section", "")

        if external_signals_text:
            doc.add_paragraph(f"분석 기간: 최근 {days}일")
            doc.add_paragraph()

            # 신호 섹션별로 파싱하여 추가
            for line in external_signals_text.split("\n"):
                if line.startswith("■"):
                    doc.add_heading(line.replace("■ ", ""), 2)
                elif line.startswith("•"):
                    doc.add_paragraph(line, style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)
        else:
            doc.add_paragraph("수집된 외부 트렌드 신호가 없습니다.")
            doc.add_paragraph()
            doc.add_paragraph(
                "외부 신호를 수집하려면:\n"
                "1. RSS 피드 자동 수집: /api/signals/fetch/rss\n"
                "2. Reddit 트렌드 수집: /api/signals/fetch/reddit\n"
                "3. 수동 입력: /api/signals/manual"
            )

    # ===== 푸터 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"© {datetime.now().year} AMORE Pacific - Confidential").italic = True

    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 파일명 생성
    filename = f"AMORE_Insight_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.post("/analyst-report")
async def export_analyst_report(request: AnalystReportRequest):
    """
    기간별 애널리스트 리포트 DOCX 생성 (8 Sections)

    Args:
        request: AnalystReportRequest with start_date, end_date, options

    Returns:
        StreamingResponse with DOCX file
    """
    logger = logging.getLogger(__name__)
    logger.info(f"Generating analyst report: {request.start_date} ~ {request.end_date}")

    # AMOREPACIFIC colors
    PACIFIC_BLUE = RGBColor(0, 28, 88)  # #001C58
    AMORE_BLUE = RGBColor(31, 87, 149)  # #1F5795
    GRAY = RGBColor(125, 125, 125)      # #7D7D7D

    try:
        # 1. Period Analysis
        analyzer = PeriodAnalyzer()
        analysis = await analyzer.analyze(request.start_date, request.end_date)

        if analysis.total_days == 0:
            raise HTTPException(
                status_code=404,
                detail=f"No data found for period {request.start_date} ~ {request.end_date}"
            )

        # 2. External Signals (optional) - Tavily 뉴스 포함 + 3-Tier 분류
        external_signals = None
        external_signals_list = []
        if request.include_external_signals:
            try:
                # 새로운 통합 함수 사용 (Tavily + RSS + Reddit + 3-Tier 분류)
                signals_result = await _get_external_signals(
                    days=analysis.total_days,
                    brands=["LANEIGE", "COSRX", "K-Beauty"],
                    include_tavily=True,
                    start_date=request.start_date,  # 3-Tier 분류용
                    end_date=request.end_date       # 3-Tier 분류용
                )
                if signals_result.get("signals"):
                    external_signals = signals_result
                    external_signals_list = signals_result.get("signals", [])
                    classified = signals_result.get("classified", {})
                    logger.info(
                        f"Collected {len(external_signals_list)} signals: "
                        f"TIER1={len(classified.get('tier1_core', []))}, "
                        f"TIER2={len(classified.get('tier2_background', []))}, "
                        f"TIER3={len(classified.get('tier3_archive', []))}"
                    )
            except Exception as e:
                logger.warning(f"External signal collection failed: {e}")

        # 3. Generate Insights
        insight_agent = PeriodInsightAgent()
        report = await insight_agent.generate_report(analysis, external_signals=external_signals)

        # 4. Generate Charts
        charts = {}
        chart_paths = {}
        if request.include_charts:
            temp_dir = tempfile.mkdtemp()
            chart_gen = ChartGenerator(output_dir=temp_dir)
            chart_paths = chart_gen.generate_all_charts(analysis)
            logger.info(f"Generated {len(chart_paths)} charts in {temp_dir}")

        # 5. Reference Tracker
        tracker = ReferenceTracker()
        tracker.auto_add_amazon_sources(
            start_date=request.start_date,
            end_date=request.end_date
        )

        # 외부 신호를 참고자료에 자동 추가 (Tavily 뉴스, RSS, Reddit 등)
        if external_signals_list:
            added_refs = tracker.add_external_signals(external_signals_list)
            logger.info(f"Added {added_refs} external signal references")

        # 6. Create DOCX Document
        doc = Document()

        # Style setup
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)

        # ===== 표지 + 목차 (Cover Page with TOC - 같은 페이지) =====
        # 제목
        title = doc.add_heading('LANEIGE Amazon US 경쟁력 분석 보고서', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = PACIFIC_BLUE
            run.font.size = Pt(24)
            run.font.bold = True

        # 구분선 (Pacific Blue)
        divider = doc.add_paragraph()
        divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
        divider_run = divider.add_run('─' * 40)
        divider_run.font.color.rgb = PACIFIC_BLUE
        divider_run.font.size = Pt(10)

        # 분석 기간 & 생성일시 (간결하게)
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run1 = meta_para.add_run(f'분석 기간: {request.start_date} ~ {request.end_date}')
        meta_run1.font.size = Pt(12)
        meta_run1.font.color.rgb = GRAY
        meta_para.add_run('\n')
        meta_run2 = meta_para.add_run(f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        meta_run2.font.size = Pt(11)
        meta_run2.font.color.rgb = GRAY

        # ===== 목차 (Table of Contents) - 같은 페이지에 배치 =====
        # add_heading 대신 add_paragraph로 여백 최소화
        toc_heading = doc.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        toc_heading.paragraph_format.space_before = Pt(18)  # 제목과 목차 사이 최소 여백
        toc_heading.paragraph_format.space_after = Pt(8)
        toc_run = toc_heading.add_run('목차')
        toc_run.font.size = Pt(14)
        toc_run.font.bold = True
        toc_run.font.color.rgb = PACIFIC_BLUE

        # 목차 항목 - 여백 최소화
        toc_items = [
            "1. Executive Summary",
            "2. LANEIGE 심층 분석",
            "3. 경쟁 환경 분석",
            "4. 시장 동향",
            "5. 외부 신호 분석",
            "6. 리스크 및 기회 요인",
            "7. 전략 제언",
            "8. 참고자료 (References)"
        ]
        for item in toc_items:
            toc_para = doc.add_paragraph()
            toc_para.paragraph_format.space_before = Pt(2)
            toc_para.paragraph_format.space_after = Pt(2)
            toc_para.paragraph_format.left_indent = Inches(0.3)
            toc_run = toc_para.add_run(f"• {item}")
            toc_run.font.size = Pt(11)
            toc_run.font.color.rgb = PACIFIC_BLUE

        doc.add_page_break()

        # ===== Section 1: Executive Summary =====
        if report.executive_summary:
            section = report.executive_summary
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            # Content
            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    else:
                        doc.add_paragraph(line)

            # Insert SoS chart if available
            if request.include_charts and 'sos_trend' in chart_paths:
                doc.add_paragraph()
                doc.add_picture(str(chart_paths['sos_trend']), width=Inches(6))

            doc.add_page_break()

        # ===== Section 2: LANEIGE 심층 분석 =====
        if report.laneige_analysis:
            section = report.laneige_analysis
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            # Content
            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    elif line.startswith('2.'):
                        doc.add_heading(line, 2)
                    else:
                        doc.add_paragraph(line)

            # Insert charts
            if request.include_charts:
                doc.add_paragraph()
                if 'sos_trend' in chart_paths:
                    doc.add_heading('일별 SoS 추이', 3)
                    doc.add_picture(str(chart_paths['sos_trend']), width=Inches(6))
                    doc.add_paragraph()

                if 'product_ranks' in chart_paths:
                    doc.add_heading('제품별 순위 변동', 3)
                    doc.add_picture(str(chart_paths['product_ranks']), width=Inches(6))

            doc.add_page_break()

        # ===== Section 3: 경쟁 환경 분석 =====
        if report.competitive_analysis:
            section = report.competitive_analysis
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    elif line.startswith('3.'):
                        doc.add_heading(line, 2)
                    else:
                        doc.add_paragraph(line)

            # Insert brand comparison chart
            if request.include_charts and 'brand_comparison' in chart_paths:
                doc.add_paragraph()
                doc.add_heading('브랜드별 점유율', 3)
                doc.add_picture(str(chart_paths['brand_comparison']), width=Inches(6))

            doc.add_page_break()

        # ===== Section 4: 시장 동향 =====
        if report.market_trends:
            section = report.market_trends
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    elif line.startswith('4.'):
                        doc.add_heading(line, 2)
                    else:
                        doc.add_paragraph(line)

            # Insert HHI chart
            if request.include_charts and 'hhi_trend' in chart_paths:
                doc.add_paragraph()
                doc.add_heading('HHI 추이', 3)
                doc.add_picture(str(chart_paths['hhi_trend']), width=Inches(6))

            doc.add_page_break()

        # ===== Section 5: 외부 신호 분석 =====
        if report.external_signals:
            section = report.external_signals
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    elif line.startswith('5.'):
                        doc.add_heading(line, 2)
                    else:
                        doc.add_paragraph(line)

            doc.add_page_break()

        # ===== Section 6: 리스크 및 기회 요인 =====
        if report.risks_opportunities:
            section = report.risks_opportunities
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    elif line.startswith('6.'):
                        doc.add_heading(line, 2)
                    else:
                        doc.add_paragraph(line)

            doc.add_page_break()

        # ===== Section 7: 전략 제언 =====
        if report.strategic_recommendations:
            section = report.strategic_recommendations
            heading = doc.add_heading(f'{section.section_id}. {section.section_title}', 1)
            heading.runs[0].font.color.rgb = PACIFIC_BLUE

            for line in section.content.split('\n'):
                if line.strip():
                    if line.startswith('■'):
                        para = doc.add_paragraph(line)
                        para.runs[0].font.bold = True
                    elif line.startswith('7.'):
                        doc.add_heading(line, 2)
                    else:
                        doc.add_paragraph(line)

            doc.add_page_break()

        # ===== Section 8: 참고자료 (References) =====
        ref_heading = doc.add_heading('8. 참고자료 (References)', 1)
        for run in ref_heading.runs:
            run.font.color.rgb = PACIFIC_BLUE
            run.font.size = Pt(14)

        # 8.1 외부 신호 출처 (본문 인용 [1], [2] 와 매칭)
        if report.references and report.references.content:
            sub_heading1 = doc.add_heading('8.1 외부 신호 출처 (External Sources)', 2)
            for run in sub_heading1.runs:
                run.font.color.rgb = AMORE_BLUE
                run.font.size = Pt(12)

            # report.references.content에서 참고자료 추출
            for line in report.references.content.split('\n'):
                if line.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.space_before = Pt(3)
                    para.paragraph_format.space_after = Pt(3)
                    # [1], [2] 등으로 시작하는 참조 번호는 볼드 처리
                    if line.strip().startswith('['):
                        # 번호 부분과 내용 분리
                        match = re.match(r'(\[\d+\])\s*(.*)', line.strip())
                        if match:
                            ref_num = match.group(1)
                            ref_content = match.group(2)
                            run_num = para.add_run(ref_num + " ")
                            run_num.font.bold = True
                            run_num.font.color.rgb = PACIFIC_BLUE
                            run_num.font.size = Pt(10)
                            run_content = para.add_run(ref_content)
                            run_content.font.size = Pt(10)
                            run_content.font.color.rgb = GRAY
                        else:
                            run = para.add_run(line)
                            run.font.size = Pt(10)
                    else:
                        run = para.add_run(line)
                        run.font.size = Pt(10)

            doc.add_paragraph()  # 간격

        # 8.2 데이터 출처 (Amazon, IR 등)
        sub_heading2 = doc.add_heading('8.2 데이터 출처 (Data Sources)', 2)
        for run in sub_heading2.runs:
            run.font.color.rgb = AMORE_BLUE
            run.font.size = Pt(12)

        # tracker에서 데이터 출처 포맷
        ref_section = tracker.format_section()
        for line in ref_section.split('\n'):
            if line.strip():
                # 하위 섹션 번호 재조정 (8.1 → 건너뛰기)
                if line.startswith('8.1') or line.startswith('8.2'):
                    continue  # 이미 외부 신호로 처리됨
                elif line.startswith('8.'):
                    # 나머지 하위 섹션
                    doc.add_heading(line, 3)
                else:
                    para = doc.add_paragraph(line)
                    para.paragraph_format.left_indent = Inches(0.25)
                    for run in para.runs:
                        run.font.size = Pt(10)

        # ===== Footer =====
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(f"© {datetime.now().year} AMORE Pacific - Confidential")
        run.italic = True
        run.font.color.rgb = GRAY

        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Generate filename
        filename = f"AMORE_Analyst_Report_{request.start_date}_{request.end_date}.docx"

        # Cleanup temp charts
        if request.include_charts and chart_paths:
            try:
                for chart_path in chart_paths.values():
                    if chart_path.exists():
                        chart_path.unlink()
                if temp_dir:
                    Path(temp_dir).rmdir()
            except Exception as e:
                logger.warning(f"Chart cleanup failed: {e}")

        logger.info(f"Analyst report generated successfully: {filename}")

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Analyst report generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@router.post("/excel")
async def export_excel(request: Request):
    """
    엑셀 데이터 내보내기 (SQLite → Excel)
    """
    try:
        # Parse request body
        body = await request.json()
        start_date = body.get("start_date")
        end_date = body.get("end_date")
        include_metrics = body.get("include_metrics", True)

        # SQLite storage 사용
        storage = get_sqlite_storage()
        await storage.initialize()

        # 출력 경로
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = f"./data/exports/AMORE_Data_{timestamp}.xlsx"

        # 엑셀 생성
        result = storage.export_to_excel(
            output_path=output_path,
            start_date=start_date,
            end_date=end_date,
            include_metrics=include_metrics
        )

        if not result.get("success"):
            raise HTTPException(status_code=500, detail=result.get("error", "Export failed"))

        file_path = Path(result["file_path"])
        if not file_path.exists():
            raise HTTPException(status_code=500, detail="Generated file not found")

        return FileResponse(
            path=str(file_path),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={file_path.name}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Excel export error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/signals/status")
async def get_signal_status():
    """
    외부 신호 API 상태 확인

    Returns:
        각 외부 데이터 소스의 설정/사용 가능 상태
    """
    return {
        "tavily": {
            "configured": bool(os.getenv("TAVILY_API_KEY")),
            "description": "뉴스 검색 API (월 1,000건 무료)",
            "docs": "https://tavily.com"
        },
        "gnews": {
            "configured": bool(os.getenv("GNEWS_API_KEY")),
            "description": "뉴스 API (일 100건 무료)",
            "docs": "https://gnews.io"
        },
        "rss_feeds": {
            "available": True,
            "count": 10,
            "sources": [
                "Allure", "Byrdie", "Cosmetics Design Asia",
                "Cosmetics Business", "Vogue Beauty", "WWD Beauty",
                "Beautyindependent", "Global Cosmetics News",
                "Happi", "CosmeticsDesign Europe"
            ]
        },
        "reddit": {
            "available": True,
            "description": "JSON API (무료, 인증 불필요)",
            "subreddits": ["AsianBeauty", "SkincareAddiction", "MakeupAddiction"]
        },
        "public_data": {
            "customs_korea": {
                "configured": bool(os.getenv("DATA_GO_KR_API_KEY")),
                "description": "관세청 수출입통계"
            },
            "mfds_korea": {
                "configured": bool(os.getenv("DATA_GO_KR_API_KEY")),
                "description": "식약처 기능성화장품 DB"
            }
        },
        "signal_classification": {
            "tier1_core": "분석 기간 ±7일 - 직접 관련 뉴스",
            "tier2_background": "±30일 또는 구조적 트렌드",
            "tier3_archive": "30일+ 이전, 장기 트렌드만 포함"
        }
    }


# ============================================================
# 비동기 작업 API (페이지 새로고침에도 다운로드 지속)
# ============================================================

class AsyncExportRequest(BaseModel):
    """비동기 내보내기 요청"""
    job_type: str  # "export_docx", "export_analyst_report", "export_excel"
    start_date: str | None = None
    end_date: str | None = None
    include_charts: bool = True
    include_external_signals: bool = True
    include_metrics: bool = True


@router.post("/async/start")
async def start_async_export(request: AsyncExportRequest):
    """
    비동기 내보내기 작업 시작

    페이지 새로고침에도 다운로드가 지속됩니다.
    작업 ID를 반환하며, /async/status/{job_id}로 진행 상태 확인 가능.

    Args:
        request: AsyncExportRequest

    Returns:
        {
            "job_id": "abc12345",
            "status": "pending",
            "message": "작업이 큐에 추가되었습니다."
        }
    """
    from src.tools.job_queue import get_job_queue

    queue = get_job_queue()
    await queue.initialize()

    # 파라미터 준비
    params = {
        "start_date": request.start_date,
        "end_date": request.end_date,
        "include_charts": request.include_charts,
        "include_external_signals": request.include_external_signals,
        "include_metrics": request.include_metrics,
    }

    # 작업 생성
    job_id = await queue.create_job(request.job_type, params)

    return {
        "job_id": job_id,
        "status": "pending",
        "message": "작업이 큐에 추가되었습니다. 진행 상태는 /api/export/async/status/{job_id}에서 확인하세요."
    }


@router.get("/async/status/{job_id}")
async def get_async_export_status(job_id: str):
    """
    비동기 내보내기 작업 상태 조회

    Args:
        job_id: 작업 ID

    Returns:
        {
            "id": "abc12345",
            "status": "running" | "completed" | "failed" | "pending",
            "progress": 50,
            "progress_message": "차트 생성 중...",
            "download_url": "/api/export/download/abc12345" (완료 시)
        }
    """
    from src.tools.job_queue import get_job_queue

    queue = get_job_queue()
    status = await queue.get_job_status(job_id)

    if not status:
        raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

    return status


@router.get("/download/{job_id}")
async def download_export_file(job_id: str):
    """
    완료된 내보내기 파일 다운로드

    Args:
        job_id: 작업 ID

    Returns:
        FileResponse with the generated file
    """
    from src.tools.job_queue import JobStatus, get_job_queue

    queue = get_job_queue()
    status = await queue.get_job_status(job_id)

    if not status:
        raise HTTPException(status_code=404, detail=f"Job not found: {job_id}")

    if status["status"] != JobStatus.COMPLETED.value:
        raise HTTPException(
            status_code=400,
            detail=f"Job not completed yet. Current status: {status['status']}"
        )

    file_path = status.get("result_file")
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="File not found or expired")

    filename = os.path.basename(file_path)

    # MIME type 결정
    if filename.endswith('.docx'):
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif filename.endswith('.xlsx'):
        media_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    else:
        media_type = "application/octet-stream"

    return FileResponse(
        path=file_path,
        media_type=media_type,
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@router.get("/async/jobs")
async def list_export_jobs(status: str | None = None, limit: int = 20):
    """
    내보내기 작업 목록 조회

    Args:
        status: 필터링할 상태 (pending, running, completed, failed)
        limit: 최대 개수

    Returns:
        작업 목록
    """
    from src.tools.job_queue import get_job_queue

    queue = get_job_queue()
    jobs = await queue.get_all_jobs(status=status, limit=limit)

    return {
        "jobs": jobs,
        "total": len(jobs)
    }
