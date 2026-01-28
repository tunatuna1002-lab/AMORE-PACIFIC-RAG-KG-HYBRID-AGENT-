"""
Export Job Handlers

비동기 내보내기 작업 핸들러들.
job_queue.py와 함께 사용하여 페이지 새로고침에도 다운로드 지속 가능.

Usage:
    from src.tools.job_queue import get_job_queue
    from src.tools.export_handlers import register_all_handlers

    queue = get_job_queue()
    await queue.initialize()
    register_all_handlers(queue)
    await queue.start_worker()
"""

import logging
import os
import tempfile
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from src.agents.period_insight_agent import PeriodInsightAgent
from src.api.dependencies import load_dashboard_data
from src.tools.chart_generator import ChartGenerator
from src.tools.job_queue import JobQueue, JobType
from src.tools.period_analyzer import PeriodAnalyzer
from src.tools.reference_tracker import ReferenceTracker
from src.tools.sqlite_storage import get_sqlite_storage

logger = logging.getLogger(__name__)

# AMOREPACIFIC colors
PACIFIC_BLUE = RGBColor(0, 28, 88)  # #001C58
AMORE_BLUE = RGBColor(31, 87, 149)  # #1F5795
GRAY = RGBColor(125, 125, 125)  # #7D7D7D


def register_all_handlers(queue: JobQueue) -> None:
    """모든 export 핸들러 등록"""
    queue.register_handler(JobType.EXPORT_DOCX.value, handle_export_docx)
    queue.register_handler(JobType.EXPORT_ANALYST_REPORT.value, handle_export_analyst_report)
    queue.register_handler(JobType.EXPORT_EXCEL.value, handle_export_excel)
    logger.info("Registered all export handlers")


async def handle_export_docx(job_id: str, params: dict, queue: JobQueue) -> str:
    """
    간단한 DOCX 리포트 생성 핸들러

    Args:
        job_id: 작업 ID
        params: {start_date, end_date, include_strategy, include_external_signals}
        queue: JobQueue 인스턴스 (진행률 업데이트용)

    Returns:
        생성된 파일 경로
    """
    await queue.update_progress(job_id, 10, "데이터 로드 중...")

    # Load data
    data = load_dashboard_data()

    await queue.update_progress(job_id, 30, "DOCX 문서 생성 중...")

    # Create document
    doc = Document()

    # Style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # Title
    title = doc.add_heading("AMORE Pacific 인사이트 리포트", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PACIFIC_BLUE

    await queue.update_progress(job_id, 50, "섹션 작성 중...")

    # 1. Summary
    doc.add_heading("1. Executive Summary", 1)
    if data.get("home", {}).get("insight_message"):
        doc.add_paragraph(data["home"]["insight_message"])
    else:
        doc.add_paragraph("데이터가 충분하지 않습니다.")

    await queue.update_progress(job_id, 70, "브랜드 분석 추가 중...")

    # 2. Brand Performance
    doc.add_heading("2. 브랜드 성과", 1)
    brand_data = data.get("brand", {})
    kpis = brand_data.get("kpis", {})
    if kpis:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "KPI"
        hdr_cells[1].text = "현재"
        hdr_cells[2].text = "변동"

        for kpi_name, kpi_data in kpis.items():
            row = table.add_row().cells
            row[0].text = kpi_name
            row[1].text = str(kpi_data.get("value", "-"))
            row[2].text = str(kpi_data.get("change", "-"))

    await queue.update_progress(job_id, 90, "파일 저장 중...")

    # Save
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"AMORE_Insight_Report_{timestamp}.docx"
    output_path = os.path.join(queue.output_dir, filename)

    doc.save(output_path)
    logger.info(f"DOCX saved: {output_path}")

    return output_path


async def handle_export_analyst_report(job_id: str, params: dict, queue: JobQueue) -> str:
    """
    애널리스트 리포트 (8 섹션) 생성 핸들러

    Args:
        job_id: 작업 ID
        params: {start_date, end_date, include_charts, include_external_signals}
        queue: JobQueue 인스턴스

    Returns:
        생성된 파일 경로
    """
    start_date = params.get("start_date")
    end_date = params.get("end_date")
    include_charts = params.get("include_charts", True)
    include_external_signals = params.get("include_external_signals", True)

    if not start_date or not end_date:
        raise ValueError("start_date and end_date are required")

    await queue.update_progress(job_id, 5, "기간 분석 중...")

    # 1. Period Analysis
    analyzer = PeriodAnalyzer()
    analysis = await analyzer.analyze(start_date, end_date)

    if analysis.total_days == 0:
        raise ValueError(f"No data found for period {start_date} ~ {end_date}")

    await queue.update_progress(job_id, 15, "데이터 처리 중...")

    # 2. External Signals
    external_signals = None
    external_signals_list = []
    if include_external_signals:
        try:
            from src.api.routes.export import _get_external_signals

            signals_result = await _get_external_signals(
                days=analysis.total_days,
                brands=["LANEIGE", "COSRX", "K-Beauty"],
                include_tavily=True,
                start_date=start_date,
                end_date=end_date,
            )
            if signals_result.get("signals"):
                external_signals = signals_result
                external_signals_list = signals_result.get("signals", [])
        except Exception as e:
            logger.warning(f"External signal collection failed: {e}")

    await queue.update_progress(job_id, 30, "AI 인사이트 생성 중...")

    # 3. Generate Insights
    insight_agent = PeriodInsightAgent()
    report = await insight_agent.generate_report(analysis, external_signals=external_signals)

    await queue.update_progress(job_id, 50, "차트 생성 중...")

    # 4. Generate Charts
    chart_paths = {}
    temp_dir = None
    if include_charts:
        temp_dir = tempfile.mkdtemp()
        chart_gen = ChartGenerator(output_dir=temp_dir)
        chart_paths = chart_gen.generate_all_charts(analysis)
        logger.info(f"Generated {len(chart_paths)} charts")

    await queue.update_progress(job_id, 60, "참고자료 추적 중...")

    # 5. Reference Tracker
    tracker = ReferenceTracker()
    tracker.auto_add_amazon_sources(start_date=start_date, end_date=end_date)
    if external_signals_list:
        tracker.add_external_signals(external_signals_list)

    await queue.update_progress(job_id, 70, "DOCX 문서 생성 중...")

    # 6. Create DOCX
    doc = Document()

    # Style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # ===== 표지 + 목차 =====
    title = doc.add_heading("LANEIGE Amazon US 경쟁력 분석 보고서", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PACIFIC_BLUE
        run.font.size = Pt(24)
        run.font.bold = True

    # 구분선
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider.add_run("─" * 40)
    divider_run.font.color.rgb = PACIFIC_BLUE

    # 메타 정보
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(f"분석 기간: {start_date} ~ {end_date}").font.color.rgb = GRAY
    meta_para.add_run("\n")
    meta_para.add_run(
        f"생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    ).font.color.rgb = GRAY

    # 목차 (제목과 충분한 간격 확보)
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    toc_heading.paragraph_format.space_before = Pt(72)  # 제목과 목차 사이 여백 증가
    toc_heading.paragraph_format.space_after = Pt(18)
    toc_run = toc_heading.add_run("목차")
    toc_run.font.size = Pt(16)
    toc_run.font.bold = True
    toc_run.font.color.rgb = PACIFIC_BLUE

    toc_items = [
        "1. Executive Summary",
        "2. LANEIGE 심층 분석",
        "3. 경쟁 환경 분석",
        "4. 시장 동향",
        "5. 외부 신호 분석",
        "6. 리스크 및 기회 요인",
        "7. 전략 제언",
        "8. 참고자료 (References)",
    ]
    for item in toc_items:
        toc_para = doc.add_paragraph()
        toc_para.paragraph_format.left_indent = Inches(0.5)
        toc_para.paragraph_format.space_before = Pt(2)
        toc_para.paragraph_format.space_after = Pt(2)
        toc_run = toc_para.add_run(f"• {item}")
        toc_run.font.size = Pt(12)
        toc_run.font.color.rgb = PACIFIC_BLUE

    doc.add_page_break()

    await queue.update_progress(job_id, 80, "섹션 작성 중...")

    # ===== Section 1: Executive Summary =====
    if report.executive_summary:
        section = report.executive_summary
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                if line.startswith("■"):
                    para = doc.add_paragraph(line)
                    para.runs[0].font.bold = True
                else:
                    doc.add_paragraph(line)

        if include_charts and "sos_trend" in chart_paths:
            doc.add_paragraph()
            doc.add_picture(str(chart_paths["sos_trend"]), width=Inches(6))

        doc.add_page_break()

    # ===== Section 2: LANEIGE 심층 분석 =====
    if report.laneige_analysis:
        section = report.laneige_analysis
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                if line.startswith("■"):
                    para = doc.add_paragraph(line)
                    para.runs[0].font.bold = True
                else:
                    doc.add_paragraph(line)

        if include_charts:
            if "sos_trend" in chart_paths:
                doc.add_heading("일별 SoS 추이", 3)
                doc.add_picture(str(chart_paths["sos_trend"]), width=Inches(6))
            if "product_ranks" in chart_paths:
                doc.add_heading("제품별 순위 변동", 3)
                doc.add_picture(str(chart_paths["product_ranks"]), width=Inches(6))

        doc.add_page_break()

    # ===== Section 3: 경쟁 환경 분석 =====
    if report.competitive_analysis:
        section = report.competitive_analysis
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        if include_charts and "brand_comparison" in chart_paths:
            doc.add_picture(str(chart_paths["brand_comparison"]), width=Inches(6))

        doc.add_page_break()

    # ===== Section 4: 시장 동향 =====
    if report.market_trends:
        section = report.market_trends
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        if include_charts and "hhi_trend" in chart_paths:
            doc.add_picture(str(chart_paths["hhi_trend"]), width=Inches(6))

        doc.add_page_break()

    await queue.update_progress(job_id, 90, "외부 신호 및 전략 섹션 작성 중...")

    # ===== Section 5: 외부 신호 분석 =====
    if report.external_signals:
        section = report.external_signals
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        doc.add_page_break()

    # ===== Section 6: 리스크 및 기회 =====
    if report.risks_opportunities:
        section = report.risks_opportunities
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        doc.add_page_break()

    # ===== Section 7: 전략 제언 =====
    if report.strategic_recommendations:
        section = report.strategic_recommendations
        heading = doc.add_heading(f"{section.section_id}. {section.section_title}", 1)
        heading.runs[0].font.color.rgb = PACIFIC_BLUE

        for line in section.content.split("\n"):
            if line.strip():
                doc.add_paragraph(line)

        doc.add_page_break()

    # ===== Section 8: 참고자료 =====
    heading = doc.add_heading("8. 참고자료 (References)", 1)
    heading.runs[0].font.color.rgb = PACIFIC_BLUE

    # 8.1 외부 자료
    doc.add_heading("8.1 외부 자료", 2)
    external_refs = tracker.get_formatted_references(source_type="external")
    if external_refs:
        for ref in external_refs.split("\n"):
            if ref.strip():
                doc.add_paragraph(ref, style="List Bullet")
    else:
        doc.add_paragraph("외부 자료 없음")

    # 8.2 데이터 소스
    doc.add_heading("8.2 데이터 소스", 2)
    data_refs = tracker.get_formatted_references(source_type="data")
    if data_refs:
        for ref in data_refs.split("\n"):
            if ref.strip():
                doc.add_paragraph(ref, style="List Bullet")

    await queue.update_progress(job_id, 95, "파일 저장 중...")

    # Save
    filename = f"AMORE_Analyst_Report_{start_date}_{end_date}.docx"
    output_path = os.path.join(queue.output_dir, filename)

    doc.save(output_path)
    logger.info(f"Analyst report saved: {output_path}")

    # Cleanup temp charts
    if temp_dir:
        import shutil

        shutil.rmtree(temp_dir, ignore_errors=True)

    return output_path


async def handle_export_excel(job_id: str, params: dict, queue: JobQueue) -> str:
    """
    Excel 데이터 내보내기 핸들러

    Args:
        job_id: 작업 ID
        params: {start_date, end_date, include_metrics}
        queue: JobQueue 인스턴스

    Returns:
        생성된 파일 경로
    """
    start_date = params.get("start_date")
    end_date = params.get("end_date")
    include_metrics = params.get("include_metrics", True)

    await queue.update_progress(job_id, 20, "SQLite 데이터 로드 중...")

    storage = get_sqlite_storage()
    await storage.initialize()

    await queue.update_progress(job_id, 50, "Excel 파일 생성 중...")

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"AMORE_Data_{timestamp}.xlsx"
    output_path = os.path.join(queue.output_dir, filename)

    storage.export_to_excel(
        output_path=output_path,
        start_date=start_date,
        end_date=end_date,
        include_metrics=include_metrics,
    )

    await queue.update_progress(job_id, 100, "완료")

    logger.info(f"Excel saved: {output_path}")
    return output_path
