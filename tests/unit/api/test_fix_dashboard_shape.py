"""
Fix D7 - dashboard JSON shape adapters
=====================================
DashboardExporter emits `products` as a dict keyed by ASIN and `categories` as a dict
keyed by category id, while several routes were written against an older shape
(products as a list with `brand`/`title`, `summary`, `brand_metrics`, `category_metrics`,
`category[*].top_products`). `src.api.dashboard_shape` adapts both shapes.
"""

from __future__ import annotations

import inspect
from io import BytesIO
from types import SimpleNamespace

import pytest

from src.api import dashboard_shape as shape
from tests.characterization._api_fixtures import exporter_shaped_dashboard

# ---------------------------------------------------------------------------
# Pure adapters
# ---------------------------------------------------------------------------


def test_products_as_list_from_exporter_dict_keyed_by_asin():
    products = shape.products_as_list(exporter_shaped_dashboard())
    assert [p["asin"] for p in products] == ["B07GFJWPDQ"]
    p = products[0]
    # exporter products are LANEIGE-only and carry `name`, not `brand`/`title`
    assert p["brand"] == "LANEIGE"
    assert p["title"] == "LANEIGE Lip Sleeping Mask"
    assert p["name"] == "LANEIGE Lip Sleeping Mask"
    assert p["product_name"] == "LANEIGE Lip Sleeping Mask"
    assert p["rank"] == 1
    assert p["category"] == "lip_care"


def test_products_as_list_passes_legacy_list_through_sorted_by_rank():
    data = {
        "products": [
            {"asin": "X", "brand": "COSRX", "title": "Snail Mucin", "rank": 3},
            {"asin": "Y", "brand": "LANEIGE", "title": "Lip Mask", "rank": 1},
        ]
    }
    out = shape.products_as_list(data)
    assert [p["asin"] for p in out] == ["Y", "X"]
    assert out[0]["name"] == "Lip Mask"
    assert out[1]["brand"] == "COSRX"  # legacy brand never overwritten


def test_products_as_list_uses_dict_key_as_asin_fallback():
    out = shape.products_as_list({"products": {"B000": {"name": "n", "rank": 2}}})
    assert out[0]["asin"] == "B000"


@pytest.mark.parametrize("data", [{}, {"products": None}, {"products": []}, {"products": "x"}])
def test_products_as_list_tolerates_missing_or_malformed(data):
    assert shape.products_as_list(data) == []


def test_categories_as_dict_accepts_dict_list_and_legacy_key():
    data = exporter_shaped_dashboard()
    assert shape.categories_as_dict(data)["lip_care"]["name"] == "Lip Care"

    as_list = {"categories": [{"category_id": "lip_care", "name": "Lip Care"}, {"id": "skin_care"}]}
    out = shape.categories_as_dict(as_list)
    assert set(out) == {"lip_care", "skin_care"}

    legacy = {"category": {"lip_care": {"top_products": []}}}
    assert shape.categories_as_dict(legacy) == {"lip_care": {"top_products": []}}

    assert shape.categories_as_dict({}) == {}


def test_summary_from_exporter_shape_derives_counts_and_avg_price():
    summary = shape.summary_from(exporter_shaped_dashboard())
    assert summary["total_products"] == 2
    assert summary["laneige_products"] == 1
    assert summary["categories_count"] == 1
    assert summary["avg_price"] == 24.0


def test_summary_from_passes_explicit_summary_through_with_defaults():
    summary = shape.summary_from({"summary": {"total_products": 5}})
    assert summary["total_products"] == 5
    assert summary["categories_count"] == 0
    assert summary["laneige_products"] == 0
    assert summary["avg_price"] == 0


def test_brand_metrics_from_exporter_competitors():
    metrics = shape.brand_metrics_from(exporter_shaped_dashboard())
    assert metrics == [
        {"brand": "LANEIGE", "sos": 50.0, "avg_rank": 1.0, "product_count": 1},
        {"brand": "Burt's Bees", "sos": 50.0, "avg_rank": 2.0, "product_count": 1},
    ]


def test_brand_metrics_from_passes_legacy_list_through():
    legacy = [{"brand": "X", "product_count": 3}]
    assert shape.brand_metrics_from({"brand_metrics": legacy}) == legacy
    assert shape.brand_metrics_from({}) == []


def test_category_metrics_from_exporter_categories():
    metrics = shape.category_metrics_from(exporter_shaped_dashboard())
    assert len(metrics) == 1
    m = metrics[0]
    assert m["category"] == "Lip Care"
    assert m["category_id"] == "lip_care"
    assert m["total_products"] == 2
    assert m["laneige_products"] == 1
    assert m["cpi"] == 150.0
    assert m["hhi"] == 0
    assert m["avg_price"] == 0


def test_category_top_products_groups_exporter_products_by_category():
    grouped = shape.category_top_products(exporter_shaped_dashboard())
    assert list(grouped) == ["lip_care"]
    assert grouped["lip_care"][0]["brand"] == "LANEIGE"
    assert grouped["lip_care"][0]["product_name"] == "LANEIGE Lip Sleeping Mask"


def test_category_top_products_prefers_legacy_top_products():
    legacy = {
        "category": {"lip_care": {"top_products": [{"brand": "LANEIGE", "product_name": "A"}]}}
    }
    assert shape.category_top_products(legacy) == {
        "lip_care": [{"brand": "LANEIGE", "product_name": "A"}]
    }


def test_ai_insights_from_falls_back_to_home_insight_message():
    insights = shape.ai_insights_from(exporter_shaped_dashboard())
    assert (
        insights["strategic_insights"][0]["content"] == "LANEIGE Lip Sleeping Mask #1 in Lip Care"
    )
    explicit = {"ai_insights": {"strategic_insights": [{"title": "t", "content": "c"}]}}
    assert shape.ai_insights_from(explicit) == explicit["ai_insights"]


def test_adapters_are_pure_and_do_not_mutate_input():
    data = exporter_shaped_dashboard()
    before = repr(data)
    for fn in (
        shape.products_as_list,
        shape.categories_as_dict,
        shape.summary_from,
        shape.brand_metrics_from,
        shape.category_metrics_from,
        shape.category_top_products,
        shape.ai_insights_from,
    ):
        fn(data)
    assert repr(data) == before
    # adapters take only the data dict
    assert all(
        len(inspect.signature(f).parameters) == 1
        for f in (shape.products_as_list, shape.summary_from)
    )


# ---------------------------------------------------------------------------
# Routes consuming the dashboard JSON
# ---------------------------------------------------------------------------


def _docx_text(content: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_export_docx_renders_laneige_from_exporter_shaped_json(
    client, dashboard_file, auth_headers, reset_rate_limits
):
    r = client.post(
        "/api/export/docx", json={"include_external_signals": False}, headers=auth_headers
    )
    assert r.status_code == 200, r.text
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    text = _docx_text(r.content)
    assert "LANEIGE" in text
    assert "LANEIGE Lip Sleeping Mask" in text  # products table (dict keyed by ASIN)
    assert "Lip Care" in text  # category metrics (dict keyed by id)
    assert "Burt's Bees" in text  # brand metrics from brand.competitors
    assert "LANEIGE 제품이 없습니다" not in text
    assert "브랜드 데이터가 없습니다" not in text
    assert "카테고리 데이터가 없습니다" not in text


def test_competitors_laneige_products_from_exporter_shaped_json(
    client, dashboard_file, reset_rate_limits, monkeypatch
):
    def _no_sqlite():
        raise RuntimeError("sqlite unavailable in test")

    monkeypatch.setattr("src.api.routes.competitors.get_sqlite_storage", _no_sqlite)

    r = client.get("/api/competitors")
    assert r.status_code == 200, r.text
    laneige = r.json()["laneige_products"]
    assert laneige, "laneige_products must be populated from data['products'] (dict keyed by ASIN)"
    assert list(laneige) == ["lip_balm"]
    product = laneige["lip_balm"][0]
    assert product["asin"] == "B07GFJWPDQ"
    assert product["category_id"] == "lip_care"
    assert product["product_type"] == "lip_balm"


def test_send_insight_report_builds_summary_from_exporter_shaped_json(
    client, dashboard_file, monkeypatch
):
    sent: dict = {}

    class FakeEmailSender:
        def is_enabled(self) -> bool:
            return True

        async def send_insight_report(self, **kwargs):
            sent.update(kwargs)
            return SimpleNamespace(success=True, sent_to=kwargs["recipients"], message="")

    class FakeStateManager:
        def get_subscription(self, email: str):
            return SimpleNamespace(email=email, verified=True)

    monkeypatch.setattr("src.tools.notifications.email_sender.EmailSender", FakeEmailSender)
    monkeypatch.setattr("src.api.routes.alerts.get_state_manager", lambda: FakeStateManager())

    r = client.post("/api/alerts/send-insight-report", json={"email": "a@b.c"})
    assert r.status_code == 200, r.text
    assert r.json()["success"] is True

    assert sent["recipients"] == ["a@b.c"]
    assert sent["avg_rank"] == 1.0
    assert sent["sos"] == 100.0  # 1 LANEIGE product out of the 1 product in the payload
    assert sent["top10_products"] == [
        {"rank": 1, "name": "LANEIGE Lip Sleeping Mask", "brand": "LANEIGE", "change": 0}
    ]
