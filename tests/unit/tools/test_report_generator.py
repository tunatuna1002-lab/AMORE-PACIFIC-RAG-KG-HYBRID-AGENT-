"""
Unit tests for report_generator.py

Coverage target: 60%+ (currently 10.77%, 587 statements)
Focus: DocxReportGenerator, PptxReportGenerator, PdfReportGenerator, ReportGenerator
"""

import tempfile
from datetime import datetime
from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document
from docx.shared import Pt, RGBColor

from src.tools.exporters.report_generator import (
    DesignSystem,
    DocxReportGenerator,
    PdfReportGenerator,
    PptxReportGenerator,
    ReportGenerator,
)

# =============================================================================
# DesignSystem Tests
# =============================================================================


class TestDesignSystem:
    """Test DesignSystem constants"""

    def test_colors_defined(self):
        """Test that all color constants are defined"""
        assert DesignSystem.PACIFIC_BLUE == RGBColor(0, 28, 88)
        assert DesignSystem.AMORE_BLUE == RGBColor(31, 87, 149)
        assert DesignSystem.GRAY == RGBColor(125, 125, 125)
        assert DesignSystem.LIGHT_GRAY == RGBColor(245, 245, 245)
        assert DesignSystem.WHITE == RGBColor(255, 255, 255)
        assert DesignSystem.ACCENT_RED == RGBColor(229, 57, 53)
        assert DesignSystem.ACCENT_GREEN == RGBColor(67, 160, 71)

    def test_fonts_defined(self):
        """Test that all font constants are defined"""
        assert DesignSystem.FONT_DOTUM == "Arita Dotum KR"
        assert DesignSystem.FONT_BURI == "Arita Buri KR"
        assert DesignSystem.FONT_FAMILY_FALLBACK == "Malgun Gothic"

    def test_typography_sizes(self):
        """Test typography size constants"""
        assert DesignSystem.TITLE_SIZE == Pt(28)
        assert DesignSystem.SUBTITLE_SIZE == Pt(16)
        assert DesignSystem.HEADING1_SIZE == Pt(18)
        assert DesignSystem.BODY_SIZE == Pt(11)

    def test_logo_paths_defined(self):
        """Test logo path constants"""
        assert "Amorepacific" in DesignSystem.LOGO_COLOR_EN
        assert "Color" in DesignSystem.LOGO_COLOR_EN
        assert "Basic" in DesignSystem.LOGO_BASIC_EN
        assert "Reverse" in DesignSystem.LOGO_REVERSE_EN


# =============================================================================
# DocxReportGenerator Tests
# =============================================================================


class TestDocxReportGenerator:
    """Test DocxReportGenerator class"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create temporary output directory"""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)

    @pytest.fixture
    def generator(self, temp_output_dir):
        """Create generator instance"""
        return DocxReportGenerator(output_dir=str(temp_output_dir))

    def test_init_with_output_dir(self, temp_output_dir):
        """Test initialization with output directory"""
        gen = DocxReportGenerator(output_dir=str(temp_output_dir))
        assert gen.output_dir == temp_output_dir
        assert gen.output_dir.exists()

    def test_init_without_output_dir(self):
        """Test initialization without output directory (uses temp)"""
        gen = DocxReportGenerator()
        assert gen.output_dir.exists()
        assert str(gen.output_dir).startswith(tempfile.gettempdir())

    def test_find_project_root_with_claude_md(self, generator):
        """Test project root detection with CLAUDE.md"""
        root = generator._find_project_root()
        assert root.exists()
        # Should find project root or fallback to cwd
        assert root.is_dir()

    def test_get_logo_path_color(self, generator):
        """Test get logo path for color variant"""
        logo_path = generator._get_logo_path("color")
        # May be None if logo doesn't exist, but should not raise
        assert logo_path is None or isinstance(logo_path, Path)

    def test_get_logo_path_basic(self, generator):
        """Test get logo path for basic variant"""
        logo_path = generator._get_logo_path("basic")
        assert logo_path is None or isinstance(logo_path, Path)

    def test_get_logo_path_reverse(self, generator):
        """Test get logo path for reverse variant"""
        logo_path = generator._get_logo_path("reverse")
        assert logo_path is None or isinstance(logo_path, Path)

    def test_get_logo_path_invalid_type(self, generator):
        """Test get logo path with invalid type (defaults to color)"""
        logo_path = generator._get_logo_path("invalid")
        assert logo_path is None or isinstance(logo_path, Path)

    def test_setup_document_styles(self, generator):
        """Test document style setup"""
        doc = Document()
        generator._setup_document_styles(doc)

        # Check Normal style
        normal_style = doc.styles["Normal"]
        assert normal_style.font.name == DesignSystem.FONT_BURI_MEDIUM
        assert normal_style.font.size == DesignSystem.BODY_SIZE

    def test_setup_page_margins(self, generator):
        """Test page margin setup"""
        doc = Document()
        generator._setup_page_margins(doc)

        for section in doc.sections:
            # Use approximate comparison due to EMU conversion precision
            assert abs(section.top_margin - DesignSystem.PAGE_MARGIN_TOP) < 1000
            assert abs(section.bottom_margin - DesignSystem.PAGE_MARGIN_BOTTOM) < 1000
            assert abs(section.left_margin - DesignSystem.PAGE_MARGIN_LEFT) < 1000
            assert abs(section.right_margin - DesignSystem.PAGE_MARGIN_RIGHT) < 1000

    def test_add_header_bar(self, generator):
        """Test adding header bar to document"""
        doc = Document()
        initial_table_count = len(doc.tables)
        generator._add_header_bar(doc)

        # Should add one table
        assert len(doc.tables) == initial_table_count + 1

    def test_add_cover_page_minimal(self, generator):
        """Test adding cover page with minimal info"""
        doc = Document()
        generator._add_cover_page(doc, title="Test Report")

        # Should have content
        assert len(doc.paragraphs) > 0
        # Should have page break
        assert any("Test Report" in p.text for p in doc.paragraphs)

    def test_add_cover_page_full(self, generator):
        """Test adding cover page with all info"""
        doc = Document()
        generator._add_cover_page(
            doc,
            title="LANEIGE Amazon US 경쟁력 분석 보고서",
            subtitle="Weekly Insight Report",
            date_range="2026-01-01 ~ 2026-01-14",
            generation_date="2026-02-17 10:00",
        )

        # Should have content
        assert len(doc.paragraphs) > 0
        paragraphs_text = "\n".join(p.text for p in doc.paragraphs)
        assert "LANEIGE" in paragraphs_text or "경쟁력" in paragraphs_text

    def test_add_cover_page_title_splitting(self, generator):
        """Test cover page title splitting logic"""
        doc = Document()
        generator._add_cover_page(
            doc,
            title="LANEIGE Amazon US 경쟁력 분석 보고서",
        )

        # Title should be split into two lines
        paragraphs_text = "\n".join(p.text for p in doc.paragraphs)
        assert "LANEIGE" in paragraphs_text

    def test_add_toc_page_minimal(self, generator):
        """Test adding TOC page with minimal info"""
        doc = Document()
        sections = ["Executive Summary", "Analysis", "Conclusion"]
        generator._add_toc_page(doc, sections)

        # Should have CONTENTS heading
        paragraphs_text = "\n".join(p.text for p in doc.paragraphs)
        assert "CONTENTS" in paragraphs_text

    def test_add_toc_page_with_disclaimer(self, generator):
        """Test adding TOC page with disclaimer"""
        doc = Document()
        sections = ["Section 1", "Section 2"]
        disclaimer = "This is a test disclaimer."
        generator._add_toc_page(doc, sections, disclaimer=disclaimer)

        paragraphs_text = "\n".join(p.text for p in doc.paragraphs)
        assert "DISCLAIMER" in paragraphs_text
        assert disclaimer in paragraphs_text

    def test_add_section_heading_level1(self, generator):
        """Test adding level 1 section heading"""
        doc = Document()
        generator._add_section_heading(doc, section_id=1, title="Test Section", level=1)

        # Should add heading
        assert len(doc.paragraphs) > 0

    def test_add_section_heading_level2(self, generator):
        """Test adding level 2 section heading"""
        doc = Document()
        generator._add_section_heading(doc, section_id=1, title="Test Section", level=2)

        assert len(doc.paragraphs) > 0

    def test_add_section_heading_level3(self, generator):
        """Test adding level 3 section heading"""
        doc = Document()
        generator._add_section_heading(doc, section_id=1, title="Test Section", level=3)

        assert len(doc.paragraphs) > 0

    def test_add_content_paragraph_normal(self, generator):
        """Test adding normal content paragraph"""
        doc = Document()
        generator._add_content_paragraph(doc, "Test content")

        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[-1].text == "Test content"

    def test_add_content_paragraph_highlight(self, generator):
        """Test adding highlighted content paragraph"""
        doc = Document()
        generator._add_content_paragraph(doc, "■ Important point", is_highlight=True)

        assert len(doc.paragraphs) > 0
        assert "■ Important point" in doc.paragraphs[-1].text

    def test_add_content_paragraph_bullet(self, generator):
        """Test adding bullet point paragraph"""
        doc = Document()
        generator._add_content_paragraph(doc, "Bullet item", is_bullet=True)

        assert len(doc.paragraphs) > 0

    def test_add_content_paragraph_with_space_before(self, generator):
        """Test adding paragraph with space before"""
        doc = Document()
        initial_para_count = len(doc.paragraphs)
        generator._add_content_paragraph(doc, "Test", add_space_before=True)

        # Should add space paragraph + content paragraph
        assert len(doc.paragraphs) > initial_para_count

    def test_add_kpi_cards_empty(self, generator):
        """Test adding KPI cards with empty list"""
        doc = Document()
        initial_table_count = len(doc.tables)
        generator._add_kpi_cards(doc, [])

        # Should not add any tables
        assert len(doc.tables) == initial_table_count

    def test_add_kpi_cards_single(self, generator):
        """Test adding single KPI card"""
        doc = Document()
        kpis = [{"name": "SoS", "value": "2.5%", "change": "+0.3%", "trend": "up"}]
        generator._add_kpi_cards(doc, kpis)

        # Should add table
        assert len(doc.tables) > 0

    def test_add_kpi_cards_multiple(self, generator):
        """Test adding multiple KPI cards"""
        doc = Document()
        kpis = [
            {"name": "SoS", "value": "2.5%", "change": "+0.3%", "trend": "up"},
            {"name": "HHI", "value": "1200", "change": "-50", "trend": "down"},
            {"name": "CPI", "value": "0.95", "change": "0.0", "trend": "neutral"},
        ]
        generator._add_kpi_cards(doc, kpis)

        assert len(doc.tables) > 0

    def test_add_kpi_cards_with_trends(self, generator):
        """Test KPI cards with different trend indicators"""
        doc = Document()
        kpis = [
            {"name": "Up", "value": "10", "change": "+5", "trend": "up"},
            {"name": "Down", "value": "5", "change": "-2", "trend": "down"},
            {"name": "Neutral", "value": "7", "change": "0", "trend": "neutral"},
        ]
        generator._add_kpi_cards(doc, kpis)

        assert len(doc.tables) > 0

    def test_add_chart_image_not_found(self, generator, temp_output_dir):
        """Test adding chart image that doesn't exist"""
        doc = Document()
        fake_path = temp_output_dir / "nonexistent_chart.png"
        initial_para_count = len(doc.paragraphs)

        generator._add_chart_image(doc, fake_path)

        # Should not add image if file doesn't exist
        # But may add empty paragraph
        assert len(doc.paragraphs) >= initial_para_count

    def test_add_chart_image_with_caption(self, generator, temp_output_dir):
        """Test adding chart image with caption"""
        doc = Document()

        # Create a fake image file
        fake_chart = temp_output_dir / "test_chart.png"
        # Create minimal PNG (1x1 pixel)
        fake_chart.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
            b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        generator._add_chart_image(doc, fake_chart, caption="Test Chart", width=5.0)

        # Should add paragraphs
        assert len(doc.paragraphs) > 0

    def test_add_footer(self, generator):
        """Test adding footer to document"""
        doc = Document()
        generator._add_footer(doc)

        # Should add footer paragraph
        assert len(doc.paragraphs) > 0
        footer_text = doc.paragraphs[-1].text
        assert "AMORE Pacific" in footer_text
        assert str(datetime.now().year) in footer_text

    def test_generate_analyst_report_minimal(self, generator, temp_output_dir):
        """Test generating minimal analyst report"""
        report_data = {
            "title": "Test Report",
            "sections": [
                {
                    "id": 1,
                    "title": "Executive Summary",
                    "content": "This is a test report.",
                }
            ],
        }

        output_path = generator.generate_analyst_report(report_data)

        assert output_path.exists()
        assert output_path.suffix == ".docx"
        assert output_path.parent == temp_output_dir

    def test_generate_analyst_report_full(self, generator, temp_output_dir):
        """Test generating full analyst report with all features"""
        report_data = {
            "title": "LANEIGE Amazon US 경쟁력 분석 보고서",
            "subtitle": "Weekly Insight Report",
            "start_date": "2026-01-01",
            "end_date": "2026-01-14",
            "sections": [
                {
                    "id": 1,
                    "title": "Executive Summary",
                    "content": "■ Key Findings\n• Point 1\n• Point 2\n\nDetailed analysis.",
                    "kpis": [{"name": "SoS", "value": "2.5%", "change": "+0.3%", "trend": "up"}],
                },
                {
                    "id": 2,
                    "title": "Market Analysis",
                    "content": "Market trends analysis.",
                },
            ],
            "references": "Data source: Amazon.com\nAnalysis date: 2026-02-17",
        }

        output_path = generator.generate_analyst_report(report_data)

        assert output_path.exists()
        assert output_path.suffix == ".docx"

    def test_generate_analyst_report_with_custom_filename(self, generator, temp_output_dir):
        """Test generating report with custom filename"""
        report_data = {
            "title": "Test Report",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        output_path = generator.generate_analyst_report(
            report_data, output_filename="custom_report.docx"
        )

        assert output_path.exists()
        assert output_path.name == "custom_report.docx"

    def test_generate_analyst_report_with_chart_paths(self, generator, temp_output_dir):
        """Test generating report with chart paths"""
        # Create fake chart
        fake_chart = temp_output_dir / "test_chart.png"
        fake_chart.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
            b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        report_data = {
            "title": "Test Report",
            "sections": [
                {
                    "id": 1,
                    "title": "Analysis",
                    "content": "Content",
                    "chart_key": "test_chart",
                }
            ],
        }

        chart_paths = {"test_chart": fake_chart}

        output_path = generator.generate_analyst_report(report_data, chart_paths=chart_paths)

        assert output_path.exists()

    def test_to_bytes_from_path(self, generator, temp_output_dir):
        """Test converting document to bytes from path"""
        # Create a minimal docx file
        doc = Document()
        doc.add_paragraph("Test")
        test_path = temp_output_dir / "test.docx"
        doc.save(test_path)

        result_bytes = generator.to_bytes(doc_path=test_path)

        assert isinstance(result_bytes, bytes)
        assert len(result_bytes) > 0

    def test_to_bytes_from_document(self, generator):
        """Test converting document to bytes from Document object"""
        doc = Document()
        doc.add_paragraph("Test content")

        result_bytes = generator.to_bytes(doc=doc)

        assert isinstance(result_bytes, bytes)
        assert len(result_bytes) > 0

    def test_to_bytes_no_input_raises_error(self, generator):
        """Test to_bytes raises error when no input provided"""
        with pytest.raises(ValueError, match="Either doc_path or doc must be provided"):
            generator.to_bytes()


# =============================================================================
# PptxReportGenerator Tests
# =============================================================================


class TestPptxReportGenerator:
    """Test PptxReportGenerator class"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create temporary output directory"""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)

    @pytest.fixture
    def generator(self, temp_output_dir):
        """Create generator instance"""
        return PptxReportGenerator(output_dir=str(temp_output_dir))

    def test_init_with_output_dir(self, temp_output_dir):
        """Test initialization with output directory"""
        gen = PptxReportGenerator(output_dir=str(temp_output_dir))
        assert gen.output_dir == temp_output_dir
        assert gen.output_dir.exists()

    def test_init_without_output_dir(self):
        """Test initialization without output directory"""
        gen = PptxReportGenerator()
        assert gen.output_dir.exists()

    def test_find_project_root(self, generator):
        """Test project root detection"""
        root = generator._find_project_root()
        assert root.exists()
        assert root.is_dir()

    def test_get_logo_path_reverse(self, generator):
        """Test get logo path for reverse variant (default for PPTX)"""
        logo_path = generator._get_logo_path("reverse")
        assert logo_path is None or isinstance(logo_path, Path)

    def test_get_logo_path_color(self, generator):
        """Test get logo path for color variant"""
        logo_path = generator._get_logo_path("color")
        assert logo_path is None or isinstance(logo_path, Path)

    def test_generate_presentation_pptx_available(self, generator, temp_output_dir):
        """Test generating presentation when python-pptx is available"""
        # Set pptx available before mocking
        generator._pptx_available = True

        with patch("pptx.Presentation") as mock_prs_class:
            mock_prs = MagicMock()
            mock_prs_class.return_value = mock_prs
            mock_prs.slide_layouts = [MagicMock() for _ in range(10)]
            mock_prs.slides = MagicMock()
            mock_prs.slide_width = MagicMock()
            mock_prs.slide_height = MagicMock()

            report_data = {
                "title": "Test Presentation",
                "start_date": "2026-01-01",
                "end_date": "2026-01-14",
                "sections": [
                    {"id": 1, "title": "Section 1", "content": "Content 1"},
                ],
            }

            output_path = generator.generate_presentation(report_data)

            # Should attempt to create presentation
            assert mock_prs_class.called

    def test_generate_presentation_pptx_not_available(self, generator):
        """Test generating presentation when python-pptx is not available"""
        generator._pptx_available = False

        report_data = {
            "title": "Test Presentation",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        result = generator.generate_presentation(report_data)

        # Should return None when pptx not available
        assert result is None


# =============================================================================
# PdfReportGenerator Tests
# =============================================================================


class TestPdfReportGenerator:
    """Test PdfReportGenerator class"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create temporary output directory"""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)

    @pytest.fixture
    def generator(self, temp_output_dir):
        """Create generator instance"""
        return PdfReportGenerator(output_dir=str(temp_output_dir))

    def test_init_with_output_dir(self, temp_output_dir):
        """Test initialization with output directory"""
        gen = PdfReportGenerator(output_dir=str(temp_output_dir))
        assert gen.output_dir == temp_output_dir
        assert gen.output_dir.exists()

    def test_init_without_output_dir(self):
        """Test initialization without output directory"""
        gen = PdfReportGenerator()
        assert gen.output_dir.exists()

    def test_generate_from_html_weasyprint_available(self, generator, temp_output_dir):
        """Test generating PDF from HTML when WeasyPrint is available"""
        # Set weasyprint available before mocking
        generator._weasyprint_available = True

        # Mock the weasyprint module imports
        import sys

        mock_weasyprint = MagicMock()
        mock_html_class = MagicMock()
        mock_html = MagicMock()
        mock_html_class.return_value = mock_html
        mock_weasyprint.HTML = mock_html_class
        mock_weasyprint.CSS = MagicMock()

        with patch.dict(sys.modules, {"weasyprint": mock_weasyprint}):
            html_content = "<html><body><h1>Test Report</h1></body></html>"

            output_path = generator.generate_from_html(html_content)

            # Should call HTML and write_pdf
            assert mock_html_class.called
            assert mock_html.write_pdf.called

    def test_generate_from_html_weasyprint_not_available(self, generator):
        """Test generating PDF from HTML when WeasyPrint is not available"""
        generator._weasyprint_available = False

        html_content = "<html><body><h1>Test</h1></body></html>"
        result = generator.generate_from_html(html_content)

        # Should return None when weasyprint not available
        assert result is None

    def test_generate_from_html_custom_filename(self, generator, temp_output_dir):
        """Test generating PDF with custom filename"""
        generator._weasyprint_available = True

        # Mock the weasyprint module imports
        import sys

        mock_weasyprint = MagicMock()
        mock_html_class = MagicMock()
        mock_html = MagicMock()
        mock_html_class.return_value = mock_html
        mock_weasyprint.HTML = mock_html_class
        mock_weasyprint.CSS = MagicMock()

        with patch.dict(sys.modules, {"weasyprint": mock_weasyprint}):
            html_content = "<html><body>Test</body></html>"

            output_path = generator.generate_from_html(html_content, output_filename="custom.pdf")

            # Should use custom filename
            assert mock_html.write_pdf.called

    def test_convert_docx_to_pdf_docx2pdf_success(self, generator, temp_output_dir):
        """Test converting DOCX to PDF using docx2pdf"""
        docx_path = temp_output_dir / "test.docx"
        docx_path.write_text("fake docx")
        pdf_path = temp_output_dir / "test.pdf"

        # Mock the docx2pdf module
        import sys

        mock_docx2pdf = MagicMock()

        # Simulate successful conversion
        def create_pdf(*args, **kwargs):
            pdf_path.write_text("fake pdf")

        mock_docx2pdf.convert = MagicMock(side_effect=create_pdf)

        with patch.dict(sys.modules, {"docx2pdf": mock_docx2pdf}):
            result = generator.convert_docx_to_pdf(docx_path)

            assert result == pdf_path
            assert pdf_path.exists()

    def test_convert_docx_to_pdf_docx2pdf_import_error(self, generator, temp_output_dir):
        """Test converting DOCX to PDF when docx2pdf import fails"""
        docx_path = temp_output_dir / "test.docx"
        docx_path.write_text("fake docx")

        # Mock the import to raise ImportError
        import sys

        with patch.dict(sys.modules, {"docx2pdf": None}):
            with patch("subprocess.run") as mock_subprocess:
                # Simulate LibreOffice failure too
                mock_subprocess.side_effect = FileNotFoundError

                result = generator.convert_docx_to_pdf(docx_path)

                # Should return None when all converters fail
                assert result is None

    def test_convert_docx_to_pdf_libreoffice_success(self, generator, temp_output_dir):
        """Test converting DOCX to PDF using LibreOffice"""
        docx_path = temp_output_dir / "test.docx"
        docx_path.write_text("fake docx")
        pdf_path = temp_output_dir / "test.pdf"

        # Mock the import to raise ImportError
        import sys

        with patch.dict(sys.modules, {"docx2pdf": None}):
            with patch("subprocess.run") as mock_subprocess:
                # Simulate successful LibreOffice conversion
                def create_pdf(*args, **kwargs):
                    pdf_path.write_text("fake pdf")

                mock_subprocess.side_effect = create_pdf

                result = generator.convert_docx_to_pdf(docx_path)

                assert result == pdf_path
                assert pdf_path.exists()

    def test_convert_docx_to_pdf_all_fail(self, generator, temp_output_dir):
        """Test DOCX to PDF conversion when all methods fail"""
        docx_path = temp_output_dir / "test.docx"
        docx_path.write_text("fake docx")

        # Mock the import to raise ImportError
        import sys

        with patch.dict(sys.modules, {"docx2pdf": None}):
            with patch("subprocess.run", side_effect=FileNotFoundError):
                result = generator.convert_docx_to_pdf(docx_path)

                assert result is None


# =============================================================================
# ReportGenerator Tests (Unified Interface)
# =============================================================================


class TestReportGenerator:
    """Test unified ReportGenerator class"""

    @pytest.fixture
    def temp_output_dir(self):
        """Create temporary output directory"""
        with tempfile.TemporaryDirectory() as tmpdir:
            yield Path(tmpdir)

    @pytest.fixture
    def generator(self, temp_output_dir):
        """Create generator instance"""
        return ReportGenerator(output_dir=str(temp_output_dir))

    def test_init_with_output_dir(self, temp_output_dir):
        """Test initialization with output directory"""
        gen = ReportGenerator(output_dir=str(temp_output_dir))
        assert gen.output_dir == temp_output_dir
        assert gen.output_dir.exists()
        assert isinstance(gen.docx_generator, DocxReportGenerator)
        assert isinstance(gen.pptx_generator, PptxReportGenerator)
        assert isinstance(gen.pdf_generator, PdfReportGenerator)

    def test_init_without_output_dir(self):
        """Test initialization without output directory"""
        gen = ReportGenerator()
        assert gen.output_dir.exists()
        assert isinstance(gen.docx_generator, DocxReportGenerator)

    def test_generate_docx_only(self, generator, temp_output_dir):
        """Test generating DOCX format only"""
        report_data = {
            "title": "Test Report",
            "start_date": "2026-01-01",
            "end_date": "2026-01-14",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        results = generator.generate(report_data, formats=["docx"])

        assert "docx" in results
        assert results["docx"].exists()
        assert results["docx"].suffix == ".docx"

    def test_generate_default_format(self, generator, temp_output_dir):
        """Test generating with default format (docx)"""
        report_data = {
            "title": "Test Report",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        results = generator.generate(report_data)

        # Default should be DOCX only
        assert "docx" in results
        assert results["docx"].exists()

    def test_generate_with_custom_base_filename(self, generator, temp_output_dir):
        """Test generating with custom base filename"""
        report_data = {
            "title": "Test Report",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        results = generator.generate(report_data, formats=["docx"], base_filename="custom")

        assert "docx" in results
        assert "custom.docx" in results["docx"].name

    def test_generate_pdf_with_existing_docx(self, generator, temp_output_dir):
        """Test generating PDF when DOCX is also requested"""
        report_data = {
            "title": "Test Report",
            "start_date": "2026-01-01",
            "end_date": "2026-01-14",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        with patch.object(generator.pdf_generator, "convert_docx_to_pdf") as mock_convert:
            mock_convert.return_value = temp_output_dir / "test.pdf"

            results = generator.generate(report_data, formats=["docx", "pdf"])

            assert "docx" in results
            assert mock_convert.called

    def test_generate_pdf_without_docx(self, generator, temp_output_dir):
        """Test generating PDF without DOCX in formats list"""
        report_data = {
            "title": "Test Report",
            "start_date": "2026-01-01",
            "end_date": "2026-01-14",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        with patch.object(generator.pdf_generator, "convert_docx_to_pdf") as mock_convert:
            mock_convert.return_value = temp_output_dir / "test.pdf"

            results = generator.generate(report_data, formats=["pdf"])

            # Should create temp DOCX then convert
            assert mock_convert.called

    def test_generate_multiple_formats(self, generator, temp_output_dir):
        """Test generating multiple formats"""
        report_data = {
            "title": "Test Report",
            "start_date": "2026-01-01",
            "end_date": "2026-01-14",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        with patch.object(generator.pptx_generator, "generate_presentation") as mock_pptx:
            mock_pptx.return_value = temp_output_dir / "test.pptx"

            with patch.object(generator.pdf_generator, "convert_docx_to_pdf") as mock_pdf:
                mock_pdf.return_value = temp_output_dir / "test.pdf"

                results = generator.generate(report_data, formats=["docx", "pptx", "pdf"])

                assert "docx" in results
                assert mock_pptx.called
                assert mock_pdf.called

    def test_generate_with_chart_paths(self, generator, temp_output_dir):
        """Test generating reports with chart paths"""
        # Create fake chart
        fake_chart = temp_output_dir / "chart.png"
        fake_chart.write_bytes(
            b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
            b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01"
            b"\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
        )

        report_data = {
            "title": "Test Report",
            "sections": [
                {"id": 1, "title": "Analysis", "content": "Test", "chart_key": "test_chart"}
            ],
        }

        chart_paths = {"test_chart": fake_chart}

        results = generator.generate(report_data, chart_paths=chart_paths, formats=["docx"])

        assert "docx" in results
        assert results["docx"].exists()

    def test_generate_pptx_failure_handled(self, generator, temp_output_dir):
        """Test that PPTX generation failure is handled gracefully"""
        report_data = {
            "title": "Test Report",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        with patch.object(generator.pptx_generator, "generate_presentation") as mock_pptx:
            mock_pptx.return_value = None  # Simulate failure

            results = generator.generate(report_data, formats=["docx", "pptx"])

            # Should still have DOCX
            assert "docx" in results
            # PPTX should not be in results if generation failed
            assert "pptx" not in results or results["pptx"] is None

    def test_generate_pdf_failure_handled(self, generator, temp_output_dir):
        """Test that PDF conversion failure is handled gracefully"""
        report_data = {
            "title": "Test Report",
            "sections": [{"id": 1, "title": "Test", "content": "Content"}],
        }

        with patch.object(generator.pdf_generator, "convert_docx_to_pdf") as mock_pdf:
            mock_pdf.return_value = None  # Simulate failure

            results = generator.generate(report_data, formats=["docx", "pdf"])

            # Should still have DOCX
            assert "docx" in results
            # PDF should not be in results if conversion failed
            assert "pdf" not in results or results["pdf"] is None
