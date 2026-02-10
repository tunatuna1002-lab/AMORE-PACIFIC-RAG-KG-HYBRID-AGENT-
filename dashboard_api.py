"""
Dashboard API Server
====================
대시보드용 FastAPI 백엔드 서버 (메인 엔트리포인트)

## 핵심 기능
- 챗봇 API (ChatGPT + RAG + Ontology 연동)
- DOCX 인사이트 리포트 생성
- 대화 메모리 지원 (세션별 TTL 기반)
- Audit Trail 로깅

## 아키텍처 흐름
```
┌─────────────────────────────────────────────────────────────────────────┐
│                           FastAPI Server                                │
│   dashboard_api.py (PORT 8001)                                          │
├─────────────────────────────────────────────────────────────────────────┤
│                                                                         │
│  /api/chat ─────────────► HybridChatbotAgent ─────────► LLM (GPT-4.1)  │
│                                   │                                     │
│                         ┌─────────┴─────────┐                          │
│                         ▼                   ▼                          │
│                  KnowledgeGraph      DocumentRetriever                  │
│                  (온톨로지)          (RAG 가이드라인)                   │
│                                                                         │
│  /api/crawl/start ────► UnifiedBrain ────► AmazonScraper               │
│                              │              (Playwright)                │
│                              ▼                                          │
│                        MetricCalculator                                 │
│                              │                                          │
│                              ▼                                          │
│                       SheetsWriter / SQLite                             │
│                                                                         │
│  /api/data ───────────► dashboard_data.json (캐시된 데이터)            │
│                                                                         │
│  /dashboard ──────────► amore_unified_dashboard_v4.html                │
└─────────────────────────────────────────────────────────────────────────┘
```

## 주요 엔드포인트
- GET  /           : 헬스체크
- GET  /api/data   : 대시보드 데이터 JSON
- POST /api/chat   : 챗봇 v1 (RAG)
- POST /api/v2/chat: 챗봇 v2 (Unified Brain)
- POST /api/v3/chat: 챗봇 v3 (Simple Chat)
- POST /api/crawl/start: 크롤링 시작 (API Key 필요)
- GET  /dashboard  : 대시보드 UI

## 환경 변수
- OPENAI_API_KEY: OpenAI API 키 (필수)
- API_KEY: 보호된 엔드포인트용 인증키
- AUTO_START_SCHEDULER: 서버 시작 시 스케줄러 자동 시작 (default: true)
"""

import asyncio
import json
import logging
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from dotenv import load_dotenv
from fastapi import Depends, HTTPException, Request
from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
from litellm import acompletion

# App Factory (미들웨어, 라우터, 정적 파일 등록 포함)
from src.api.app_factory import create_app

# 공통 의존성 (인증, 세션, 헬퍼 등)
from src.api.dependencies import (
    add_to_memory,
    build_data_context,
    conversation_memory,
    generate_dynamic_suggestions,
    get_base_url,
    get_conversation_history,
    get_rag_context,
    limiter,
    load_dashboard_data,
    log_chat_interaction,
    rag_router,
    verify_api_key,
)

# Pydantic 모델
from src.api.models import (
    AlertSendRequest,
    AlertSettingsRequest,
    BrainChatRequest,
    BrainChatResponse,
    ChatRequest,
    ChatResponse,
    DealsRequest,
    DealsResponse,
    ExportRequest,
    SubscribeRequest,
    UpdateAlertSettingsRequest,
)

# Core 모듈 (직접 의존)
from src.core.brain import get_initialized_brain
from src.core.crawl_manager import get_crawl_manager
from src.core.state_manager import get_state_manager
from src.rag.router import QueryType
from src.tools.storage.sqlite_storage import get_sqlite_storage

load_dotenv()

logger = logging.getLogger(__name__)

# App 생성 (app_factory에서 미들웨어, 라우터, 정적 파일 등록 완료)
app = create_app()


# 글로벌 예외 핸들러 - 에러 발생 시 Telegram 알림
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    """모든 예외를 잡아서 Telegram 알림 전송"""

    error_detail = f"{type(exc).__name__}: {str(exc)[:200]}"
    endpoint = f"{request.method} {request.url.path}"

    # 로깅
    logger.error(f"Unhandled exception at {endpoint}: {error_detail}")

    # Telegram 알림 (비동기, 실패해도 무시)
    try:
        from src.tools.notifications.telegram_bot import notify_error

        asyncio.create_task(notify_error(exc, context=f"API: {endpoint}"))
    except Exception:
        pass  # Telegram 알림 실패는 무시

    # 클라이언트에게는 일반 에러 응답
    return JSONResponse(
        status_code=500,
        content={"error": "Internal server error", "detail": error_detail},
    )


# ============= 서버 시작 시 자동 스케줄러 =============

# Railway 배포 시 healthcheck 타임아웃 방지: 기본값 false
# 로컬 개발 시 AUTO_START_SCHEDULER=true 로 설정하면 스케줄러 자동 시작
AUTO_START_SCHEDULER = os.getenv("AUTO_START_SCHEDULER", "false").lower() == "true"


@app.on_event("startup")
async def startup_event():
    """서버 시작 시 설정 검증, 자동 스케줄러 시작 및 즉시 크롤링 체크

    ⚠️ 중요: 크롤링은 백그라운드에서 실행하여 healthcheck 타임아웃 방지
    """
    # 0. 설정 검증 (필수 설정 누락 시 경고, 서버는 계속 시작)
    try:
        from src.infrastructure.config.config_manager import AppConfig

        config = AppConfig.from_env_validated(fail_fast=False)
        logging.info(
            f"설정 검증 완료 (port={config.port}, scheduler={config.auto_start_scheduler})"
        )
    except Exception as e:
        logging.error(f"설정 검증 중 오류: {e}")

    # 1. 크롤링 필요 여부 체크 후 백그라운드 실행 (비블로킹)
    try:
        crawl_manager = await get_crawl_manager()
        if crawl_manager.needs_crawl():
            logging.info(
                f"서버 시작: 오늘({crawl_manager.get_kst_today()}) 데이터 없음 → 크롤링 백그라운드 시작"
            )
            # ⚠️ await 대신 create_task로 백그라운드 실행 (healthcheck 블로킹 방지)
            asyncio.create_task(crawl_manager.start_crawl())
        else:
            logging.info(
                f"서버 시작: 오늘 데이터 있음 또는 크롤링 중 (data_date={crawl_manager.get_data_date()})"
            )
    except Exception as e:
        logging.error(f"서버 시작 크롤링 체크 실패: {e}")

    # 2. 자율 스케줄러 시작 (매일 06:00 정기 크롤링용)
    if AUTO_START_SCHEDULER:
        try:
            brain = await get_initialized_brain()
            await brain.start_scheduler()
            logging.info("자율 스케줄러 자동 시작 완료 (매일 한국시간 06:00 크롤링)")
        except Exception as e:
            logging.error(f"자율 스케줄러 자동 시작 실패: {e}")

    # 3. Export Job Queue Worker 시작 (비동기 내보내기용)
    try:
        from src.tools.exporters.export_handlers import register_all_handlers
        from src.tools.utilities.job_queue import get_job_queue

        queue = get_job_queue()
        await queue.initialize()
        register_all_handlers(queue)
        await queue.start_worker()
        logging.info("Export Job Queue Worker 시작 완료")
    except Exception as e:
        logging.error(f"Export Job Queue Worker 시작 실패: {e}")

    # 4. Telegram Admin Bot 알림 (서버 시작)
    try:
        from src.tools.notifications.telegram_bot import get_bot

        bot = get_bot()
        if bot.is_enabled():
            await bot.send_alert("🚀 서버 시작됨", level="info")
            logging.info("Telegram Admin Bot 활성화됨")
    except Exception as e:
        logging.debug(f"Telegram Bot 알림 실패 (무시): {e}")


# ============= API Endpoints (helpers imported from src.api.dependencies) =============


@app.get("/api/data")
async def get_data():
    """대시보드 데이터 조회"""
    data = load_dashboard_data()
    if not data:
        raise HTTPException(status_code=404, detail="Dashboard data not found")
    return data


@app.post("/api/chat", response_model=ChatResponse, dependencies=[Depends(verify_api_key)])
@limiter.limit("10/minute")  # 분당 10회 제한 (보안 강화)
async def chat(request: Request, body: ChatRequest):
    """
    ChatGPT + RAG + Ontology 통합 챗봇 API

    1. 질문 분석 (RAGRouter)
    2. 엔티티 추출 (Ontology 기반)
    3. 관련 문서 검색 (RAG)
    4. 데이터 컨텍스트 구성
    5. 대화 기록 참조
    6. LLM 응답 생성
    7. Audit Trail 로깅
    """
    import time

    start_time = time.time()

    message = body.message.strip()
    session_id = body.session_id or "default"

    if not message:
        raise HTTPException(status_code=400, detail="Message is required")

    # 1. 질문 분류 (RAGRouter 사용)
    route_result = rag_router.route(message)
    query_type = route_result["query_type"]
    confidence = route_result["confidence"]

    # 2. 엔티티 추출 (Ontology 기반)
    entities = rag_router.extract_entities(message)

    # 3. 명확화 필요 여부 확인
    clarification = rag_router.needs_clarification(route_result, entities)
    if clarification and confidence < 0.5:
        # 명확화 요청
        add_to_memory(session_id, "user", message)
        add_to_memory(session_id, "assistant", clarification)

        return ChatResponse(
            response=clarification,
            query_type=query_type.value if hasattr(query_type, "value") else str(query_type),
            confidence=confidence,
            sources=[],
            suggestions=[
                "예, 전체 브랜드 분석해주세요",
                "LANEIGE만 분석해주세요",
                "Lip Care 카테고리만",
            ],
            entities=entities,
        )

    # 4. RAG 컨텍스트 검색
    rag_context, sources = await get_rag_context(message, query_type)

    # 5. 데이터 로드 및 컨텍스트 구성
    data = load_dashboard_data()
    data_context = build_data_context(data, query_type, entities)

    # 6. 대화 기록 조회
    conversation_history = get_conversation_history(session_id)

    # 7. 시스템 프롬프트 구성
    system_prompt = """당신은 AMORE Pacific의 LANEIGE 브랜드 Amazon 분석 전문가입니다.

역할:
- Amazon US 베스트셀러 데이터를 분석하여 인사이트 제공
- LANEIGE 브랜드의 시장 포지션 분석
- 경쟁사 대비 전략적 권고 제공
- 지표 정의 및 해석 가이드 제공

Ontology 엔티티 이해:
- Brand: 브랜드 정보 (LANEIGE, 경쟁사 등)
- Product: 제품 정보 (ASIN, 순위, 평점, 가격 등)
- Category: 카테고리 (Lip Care, Skin Care 등)
- BrandMetrics: SoS, 평균순위, 제품수 등
- ProductMetrics: 순위변동성, 연속체류일, 평점추세 등
- MarketMetrics: HHI(시장집중도), 교체율 등

응답 가이드라인:
1. 데이터에 기반한 구체적인 수치 인용
2. RAG 문서의 정의/해석 기준 활용
3. 이전 대화 맥락 고려
4. 간결하고 액션 가능한 인사이트 제공
5. 불확실한 경우 명확히 밝힐 것
6. 단정적 표현 피하기
7. 한국어로 응답

질문 유형별 응답 스타일:
- 정의(DEFINITION): 지표의 정의, 산출식, 의미를 설명
- 해석(INTERPRETATION): 수치의 의미, 좋고 나쁨의 기준 설명
- 조합(COMBINATION): 여러 지표를 함께 해석, 시나리오별 액션 제안
- 데이터조회(DATA_QUERY): 현재 수치와 변동 현황 안내
- 분석(ANALYSIS): 종합 분석과 전략적 권고 제공
"""

    # 8. 사용자 프롬프트 구성
    user_prompt = f"""## 사용자 질문
{message}

## 질문 유형
{query_type.value if hasattr(query_type, 'value') else str(query_type)} (신뢰도: {confidence:.1%})

## 추출된 엔티티
- 브랜드: {', '.join(entities.get('brands', [])) or '없음'}
- 카테고리: {', '.join(entities.get('categories', [])) or '없음'}
- 지표: {', '.join(entities.get('indicators', [])) or '없음'}
- 기간: {entities.get('time_range') or '없음'}

## RAG 참조 문서
{rag_context if rag_context else '관련 문서 없음'}

## 현재 데이터
{data_context}

## 이전 대화
{conversation_history if conversation_history else '이전 대화 없음'}

위 정보를 바탕으로 질문에 답변해주세요.
- 질문 유형에 맞는 응답 스타일을 사용하세요.
- RAG 문서에 관련 정의/해석이 있으면 인용하세요.
- 이전 대화 맥락이 있으면 고려하세요.
"""

    try:
        response = await acompletion(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=1000,
        )

        answer = response.choices[0].message.content

        # 9. 대화 메모리에 저장
        add_to_memory(session_id, "user", message)
        add_to_memory(session_id, "assistant", answer)

        # 10. 동적 후속 질문 제안 (v2 - 개선 버전)
        suggestions = generate_dynamic_suggestions(query_type, entities, answer, message)

        # 11. Audit Trail 로깅
        response_time_ms = (time.time() - start_time) * 1000
        log_chat_interaction(
            session_id=session_id,
            user_query=message,
            ai_response=answer,
            query_type=query_type.value if hasattr(query_type, "value") else str(query_type),
            confidence=confidence,
            entities=entities,
            sources=sources,
            response_time_ms=response_time_ms,
        )

        return ChatResponse(
            response=answer,
            query_type=query_type.value if hasattr(query_type, "value") else str(query_type),
            confidence=confidence,
            sources=sources,
            suggestions=suggestions,
            entities=entities,
        )

    except Exception as e:
        logger.error(f"LLM Error: {e}")

        # Fallback 응답
        fallback = route_result.get("fallback_message") or rag_router.get_fallback_response(
            "unknown"
        )

        # 데이터 기반 기본 응답 추가
        if data and query_type == QueryType.DATA_QUERY:
            brand_kpis = data.get("brand", {}).get("kpis", {})
            fallback = f"""현재 LANEIGE 현황:
- SoS: {brand_kpis.get('sos', 0)}%
- Top 10 제품: {brand_kpis.get('top10_count', 0)}개
- 평균 순위: {brand_kpis.get('avg_rank', 0)}위

(상세 분석을 위해 잠시 후 다시 시도해주세요)"""

        # Fallback 응답도 Audit Trail 기록
        response_time_ms = (time.time() - start_time) * 1000
        log_chat_interaction(
            session_id=session_id,
            user_query=message,
            ai_response=f"[ERROR] {str(e)[:100]} | Fallback: {fallback[:200]}",
            query_type=query_type.value if hasattr(query_type, "value") else str(query_type),
            confidence=0.0,
            entities=entities,
            sources=["fallback"],
            response_time_ms=response_time_ms,
        )

        return ChatResponse(
            response=fallback,
            query_type=query_type.value if hasattr(query_type, "value") else str(query_type),
            confidence=0.0,
            sources=[],
            suggestions=["다시 질문해주세요", "SoS가 뭔가요?", "현재 순위 알려주세요"],
            entities=entities,
        )


@app.delete("/api/chat/memory/{session_id}")
async def clear_memory(session_id: str):
    """세션 대화 기록 초기화"""
    if session_id in conversation_memory:
        del conversation_memory[session_id]
    return {"status": "ok", "message": f"Session {session_id} memory cleared"}


@app.get("/api/crawl/status")
async def get_crawl_status():
    """
    크롤링 상태 조회

    Returns:
        - status: idle/running/completed/failed
        - date: 크롤링 대상 날짜
        - progress: 진행률 (0-100)
        - data_date: 현재 데이터 날짜
        - needs_crawl: 크롤링 필요 여부
    """
    crawl_manager = await get_crawl_manager()
    return {
        **crawl_manager.state.to_dict(),
        "data_date": crawl_manager.get_data_date(),
        "needs_crawl": crawl_manager.needs_crawl(),
        "is_today_available": crawl_manager.is_today_data_available(),
        "status_message": crawl_manager.get_status_message(),
    }


@app.post("/api/crawl/start", dependencies=[Depends(verify_api_key)])
async def start_crawl():
    """
    수동으로 크롤링 시작 (API Key 필요)

    Returns:
        - started: 크롤링 시작 여부
        - message: 상태 메시지
    """
    crawl_manager = await get_crawl_manager()

    if crawl_manager.is_crawling():
        return {
            "started": False,
            "message": "크롤링이 이미 진행 중입니다.",
            "status": crawl_manager.state.to_dict(),
        }

    if crawl_manager.is_today_data_available():
        return {
            "started": False,
            "message": "오늘 데이터가 이미 존재합니다.",
            "status": crawl_manager.state.to_dict(),
        }

    started = await crawl_manager.start_crawl()
    return {
        "started": started,
        "message": "크롤링을 시작했습니다." if started else "크롤링 시작 실패",
        "status": crawl_manager.state.to_dict(),
    }


# ============= Historical Data API =============

from datetime import UTC, timedelta

from src.api.dependencies import get_sheets_writer


@app.get("/api/historical")
async def get_historical_data(
    start_date: str, end_date: str, category_id: str | None = None, brand: str | None = "LANEIGE"
):
    """
    히스토리컬 데이터 조회 (SQLite 우선, Google Sheets fallback)

    Args:
        start_date: 시작 날짜 (YYYY-MM-DD)
        end_date: 종료 날짜 (YYYY-MM-DD)
        category_id: 카테고리 필터 (선택)
        brand: 브랜드 필터 (기본값: LANEIGE)

    Returns:
        - data: 날짜별 지표 데이터
        - sos_history: SoS 추이 데이터
        - raw_data: 순위 추이 데이터
    """
    try:
        records = []
        data_source = None

        # 1차: SQLite에서 조회 (빠름)
        try:
            sqlite = get_sqlite_storage()
            await sqlite.initialize()
            records = await sqlite.get_raw_data(
                start_date=start_date,
                end_date=end_date,
                category_id=category_id,
                limit=50000,  # 충분히 큰 limit
            )
            if records:
                data_source = "sqlite"
                logging.info(
                    f"Historical: loaded {len(records)} records from SQLite ({start_date} ~ {end_date})"
                )
        except Exception as sqlite_err:
            logging.warning(f"Historical: SQLite 조회 실패: {sqlite_err}")

        # 2차: SQLite 실패/빈 결과 시 Google Sheets fallback
        if not records:
            try:
                sheets_writer = get_sheets_writer()
                if not sheets_writer._initialized:
                    await sheets_writer.initialize()
                records = await sheets_writer.get_raw_data(
                    start_date=start_date, end_date=end_date, category_id=category_id
                )
                if records:
                    data_source = "sheets"
                    logging.info(
                        f"Historical: loaded {len(records)} records from Sheets ({start_date} ~ {end_date})"
                    )
            except Exception as sheets_err:
                logging.warning(f"Historical: Google Sheets 조회 실패: {sheets_err}")

        if not records:
            # 모든 소스에서 데이터 없음 - 로컬 JSON 파일에서 시도
            return await _get_historical_from_local(start_date, end_date, brand)

        # 날짜 범위 계산
        start_dt = datetime.strptime(start_date, "%Y-%m-%d")
        end_dt = datetime.strptime(end_date, "%Y-%m-%d")
        days = (end_dt - start_dt).days + 1

        # 날짜별 데이터 집계 (특정 브랜드 필터링)
        daily_data = {}
        brand_lower = brand.lower() if brand else ""
        for record in records:
            snapshot_date = record.get("snapshot_date", "")
            if not snapshot_date or snapshot_date < start_date or snapshot_date > end_date:
                continue

            # 특정 브랜드 필터링 (SoS 추이 계산용)
            record_brand = record.get("brand", "")
            if brand_lower and record_brand.lower() != brand_lower:
                continue

            if snapshot_date not in daily_data:
                daily_data[snapshot_date] = {
                    "date": snapshot_date,
                    "products": [],
                    "total_count": 0,
                    "top10_count": 0,
                }

            rank = int(record.get("rank", 0)) if record.get("rank") else 0
            daily_data[snapshot_date]["products"].append(
                {
                    "asin": record.get("asin", ""),
                    "product_name": record.get("product_name", ""),
                    "brand": record_brand,
                    "rank": rank,
                    "price": record.get("price", ""),
                    "rating": record.get("rating", ""),
                }
            )
            daily_data[snapshot_date]["total_count"] += 1
            if rank <= 10:
                daily_data[snapshot_date]["top10_count"] += 1

        # SoS 추이 계산 (Top 100 기준, 해당 브랜드 기준)
        sos_history = []
        raw_data = []
        for date_str in sorted(daily_data.keys()):
            day_data = daily_data[date_str]
            products = day_data["products"]

            # SoS = (브랜드 제품 수 / 100) * 100
            sos = round(len(products) / 100 * 100, 1) if products else 0
            sos_history.append(
                {
                    "date": date_str,
                    "sos": sos,
                    "product_count": len(products),
                    "top10_count": day_data["top10_count"],
                }
            )

            # 평균 순위 (있는 경우)
            if products:
                avg_rank = round(sum(p["rank"] for p in products) / len(products), 1)
                raw_data.append(
                    {
                        "date": date_str,
                        "rank": avg_rank,
                        "best_rank": min(p["rank"] for p in products),
                        "worst_rank": max(p["rank"] for p in products),
                    }
                )

        # available_dates 계산
        available_dates = sorted(daily_data.keys())

        # brand_metrics 계산 (전체 기간 통합 - 모든 브랜드 포함)
        brand_metrics = await _calculate_brand_metrics_for_period(records, daily_data, brand)

        # rank_history 생성 (Product View 차트용)
        # 형식: { "2026-01-14": { "products": [{ "name": "...", "rank": 5, "price": 21.5 }, ...] } }
        rank_history = {}
        for record in records:
            snapshot_date = record.get("snapshot_date", "")
            if not snapshot_date or snapshot_date < start_date or snapshot_date > end_date:
                continue

            if snapshot_date not in rank_history:
                rank_history[snapshot_date] = {"products": []}

            rank = int(record.get("rank", 0)) if record.get("rank") else 0
            price_val = record.get("price", 0)
            try:
                price = float(str(price_val).replace("$", "").replace(",", "")) if price_val else 0
            except (ValueError, TypeError):
                price = 0

            rank_history[snapshot_date]["products"].append(
                {
                    "name": record.get("product_name", ""),
                    "product_name": record.get("product_name", ""),
                    "brand": record.get("brand", ""),
                    "asin": record.get("asin", ""),
                    "rank": rank,
                    "price": price,
                    "rating": record.get("rating", ""),
                    "discount_percent": record.get("discount_percent", 0),
                }
            )

        # 전체 데이터의 사용 가능한 날짜 범위 조회 (SQLite에서)
        available_date_range = {"min": None, "max": None}
        try:
            sqlite = get_sqlite_storage()
            stats = sqlite.get_stats()
            if "date_range" in stats:
                available_date_range = stats["date_range"]
        except Exception:
            pass

        return {
            "success": True,
            "available_dates": available_dates,
            "available_date_range": available_date_range,
            "data_source": data_source,
            "brand_metrics": brand_metrics,
            "rank_history": rank_history,
            "data": {
                "sos_history": sos_history,
                "raw_data": raw_data,
                "daily_data": list(daily_data.values()),
                "period": {"start": start_date, "end": end_date, "days": days},
                "brand": brand,
            },
        }

    except Exception as e:
        logging.error(f"Historical data error: {e}")
        # 폴백: 로컬 데이터에서 시도
        return await _get_historical_from_local(start_date, end_date, brand)


async def _calculate_brand_metrics_for_period(
    records: list[dict], daily_data: dict, target_brand: str
) -> list[dict]:
    """
    기간 내 모든 브랜드의 메트릭 계산 (SoS × Avg Rank 차트용)

    Note:
        기간 조회 시 동일 ASIN이 여러 날짜에 중복 등장하므로,
        ASIN 기준 유니크 카운트를 적용하여 정확한 제품 수 계산

    Returns:
        브랜드별 SoS, 평균 순위, 제품 수 등
    """
    # 전체 제품 데이터 집계 (모든 브랜드)
    brand_data = {}
    brand_unique_asins: dict[str, set] = {}  # ASIN 중복 제거용

    for record in records:
        brand_name = record.get("brand", "Unknown")
        asin = record.get("asin", "")
        rank = int(record.get("rank", 0)) if record.get("rank") else 0

        # Unknown 브랜드 및 빈 브랜드 제외 (대시보드에서 의미 없음)
        if not brand_name or brand_name.lower() == "unknown" or rank == 0:
            continue

        if brand_name not in brand_data:
            brand_data[brand_name] = {
                "brand": brand_name,
                "ranks": [],
                "prices": [],
                "product_count": 0,
            }
            brand_unique_asins[brand_name] = set()

        # 순위는 모든 레코드에서 수집 (평균 계산용)
        brand_data[brand_name]["ranks"].append(rank)

        # 가격 수집 (유효한 USD 가격 범위만)
        price = record.get("price")
        if price is not None:
            try:
                price_val = float(price)
                if 0.5 <= price_val <= 500:
                    brand_data[brand_name]["prices"].append(price_val)
            except (ValueError, TypeError):
                pass

        # 제품 수는 ASIN 기준 유니크 카운트 (중복 제거)
        if asin and asin not in brand_unique_asins[brand_name]:
            brand_unique_asins[brand_name].add(asin)
            brand_data[brand_name]["product_count"] += 1
        elif not asin:
            # ASIN이 없는 경우 기존 방식으로 카운트 (폴백)
            brand_data[brand_name]["product_count"] += 1

    # 총 유니크 제품 수 (모든 브랜드 - Unknown 제외 후)
    total_products = sum(b["product_count"] for b in brand_data.values())

    # 메트릭 계산
    brand_metrics = []
    for brand_name, data in brand_data.items():
        if not data["ranks"]:
            continue

        sos = round(data["product_count"] / max(total_products, 100) * 100, 2)
        avg_rank = round(sum(data["ranks"]) / len(data["ranks"]), 1)

        # 평균 가격 계산
        prices = data.get("prices", [])
        avg_price = round(sum(prices) / len(prices), 2) if prices else None

        # 버블 크기: 제품 수 기반 (최소 5, 최대 25)
        bubble_size = max(5, min(25, data["product_count"] * 2))

        is_laneige = target_brand.upper() in brand_name.upper()

        brand_metrics.append(
            {
                "brand": brand_name,
                "sos": sos,
                "avg_rank": avg_rank,
                "product_count": data["product_count"],
                "avg_price": avg_price,
                "bubble_size": bubble_size,
                "is_laneige": is_laneige,
            }
        )

    # SoS 기준 내림차순 정렬
    brand_metrics.sort(key=lambda x: x["sos"], reverse=True)

    # 상위 10개 추출
    top_10 = brand_metrics[:10]

    # LANEIGE가 top_10에 포함되어 있는지 확인 (brand_data에 있는지가 아니라 top_10에 있는지!)
    laneige_in_top10 = any(b.get("is_laneige") for b in top_10)

    # LANEIGE가 top_10에 없으면 추가 (데이터가 존재할 경우)
    if not laneige_in_top10 and target_brand:
        # brand_data에서 LANEIGE 찾기 (대소문자 변형 모두 시도)
        laneige_data = None
        for key in [
            target_brand,
            target_brand.upper(),
            target_brand.lower(),
            target_brand.capitalize(),
        ]:
            if key in brand_data:
                laneige_data = brand_data[key]
                break

        if laneige_data and laneige_data["ranks"]:
            sos = round(laneige_data["product_count"] / max(total_products, 100) * 100, 2)
            avg_rank = round(sum(laneige_data["ranks"]) / len(laneige_data["ranks"]), 1)
            l_prices = laneige_data.get("prices", [])
            l_avg_price = round(sum(l_prices) / len(l_prices), 2) if l_prices else None
            bubble_size = max(5, min(25, laneige_data["product_count"] * 2))
            top_10.append(
                {
                    "brand": target_brand,
                    "sos": sos,
                    "avg_rank": avg_rank,
                    "product_count": laneige_data["product_count"],
                    "avg_price": l_avg_price,
                    "bubble_size": bubble_size,
                    "is_laneige": True,
                }
            )
            # 다시 정렬 후 상위 11개 유지 (LANEIGE 포함 보장)
            top_10.sort(key=lambda x: x["sos"], reverse=True)

    # Summer Fridays 특별 처리 (고객 요청 tracked competitor)
    # top_10에 없으면 강제 추가
    TRACKED_COMPETITORS = ["Summer Fridays"]
    for tracked_brand in TRACKED_COMPETITORS:
        tracked_in_top = any(b.get("brand") == tracked_brand for b in top_10)
        if not tracked_in_top and tracked_brand in brand_data:
            tracked_data = brand_data[tracked_brand]
            if tracked_data["ranks"]:
                sos = round(tracked_data["product_count"] / max(total_products, 100) * 100, 2)
                avg_rank = round(sum(tracked_data["ranks"]) / len(tracked_data["ranks"]), 1)
                t_prices = tracked_data.get("prices", [])
                t_avg_price = round(sum(t_prices) / len(t_prices), 2) if t_prices else None
                bubble_size = max(5, min(25, tracked_data["product_count"] * 2))
                top_10.append(
                    {
                        "brand": tracked_brand,
                        "sos": sos,
                        "avg_rank": avg_rank,
                        "product_count": tracked_data["product_count"],
                        "avg_price": t_avg_price,
                        "bubble_size": bubble_size,
                        "is_laneige": False,
                        "is_tracked": True,  # tracked competitor 표시
                    }
                )
        # 데이터가 없어도 placeholder 추가 (UI에서 "-" 대신 "데이터 없음" 표시 가능)
        elif not tracked_in_top:
            top_10.append(
                {
                    "brand": tracked_brand,
                    "sos": 0,
                    "avg_rank": None,
                    "product_count": 0,
                    "bubble_size": 5,
                    "is_laneige": False,
                    "is_tracked": True,
                    "no_data": True,  # 해당 기간 데이터 없음 표시
                }
            )

    # 최종 정렬 (SoS 내림차순, tracked는 하단에 유지)
    top_10.sort(key=lambda x: (not x.get("is_tracked", False), x["sos"]), reverse=True)

    return top_10


def _get_brand_metrics_from_dashboard(dashboard_data: dict | None, target_brand: str) -> list[dict]:
    """
    대시보드 데이터에서 브랜드 메트릭 추출 (로컬 폴백용)
    """
    if not dashboard_data:
        return []

    # 대시보드의 brand_matrix 데이터 사용
    brand_matrix = dashboard_data.get("charts", {}).get("brand_matrix", [])
    if brand_matrix:
        return brand_matrix

    # 경쟁사 데이터에서 생성
    competitors = dashboard_data.get("brand", {}).get("competitors", [])
    if not competitors:
        return []

    brand_metrics = []
    for comp in competitors:
        brand_metrics.append(
            {
                "brand": comp.get("brand", "Unknown"),
                "sos": comp.get("sos", 0),
                "avg_rank": comp.get("avg_rank", 50),
                "product_count": comp.get("product_count", 0),
                "bubble_size": max(5, min(25, comp.get("product_count", 0) * 2)),
                "is_laneige": target_brand.upper() in comp.get("brand", "").upper(),
            }
        )

    return brand_metrics


async def _get_historical_from_local(
    start_date: str, end_date: str, brand: str = "LANEIGE"
) -> dict[str, Any]:
    """
    로컬 JSON 파일에서 히스토리컬 데이터 조회 (폴백)

    data/ 폴더의 날짜별 JSON 파일이나 dashboard_data.json의 히스토리 데이터 활용
    """
    try:
        # 메인 대시보드 데이터 로드
        data = load_dashboard_data()
        sos_history = []
        raw_data = []

        # 1. 대시보드 데이터에서 현재 SoS/순위 정보 추출
        if data:
            brand_kpis = data.get("brand", {}).get("kpis", {})
            current_sos = brand_kpis.get("sos", 0)
            data_date = data.get("metadata", {}).get(
                "data_date", datetime.now().strftime("%Y-%m-%d")
            )

            # 현재 날짜가 요청 범위에 포함되면 추가
            if start_date <= data_date <= end_date:
                sos_history.append(
                    {
                        "date": data_date,
                        "sos": current_sos,
                        "product_count": brand_kpis.get("product_count", 0),
                        "top10_count": brand_kpis.get("top10_count", 0),
                    }
                )

                avg_rank = brand_kpis.get("avg_rank", 0)
                if avg_rank:
                    raw_data.append(
                        {
                            "date": data_date,
                            "rank": avg_rank,
                            "best_rank": brand_kpis.get("best_rank", avg_rank),
                            "worst_rank": brand_kpis.get("worst_rank", avg_rank),
                        }
                    )

        # 2. latest_crawl_result.json에서 데이터 추출
        latest_crawl_path = Path("./data/latest_crawl_result.json")
        if latest_crawl_path.exists():
            try:
                with open(latest_crawl_path, encoding="utf-8") as f:
                    crawl_data = json.load(f)

                # 모든 카테고리에서 브랜드 제품 찾기
                brand_products = []
                crawl_date = None

                for _cat_id, cat_data in crawl_data.get("categories", {}).items():
                    for product in cat_data.get("products", []):
                        product_brand = product.get("brand", "")
                        product_name = product.get("product_name", "")

                        # 브랜드 매칭 (대소문자 무시, 부분 매칭)
                        if (
                            brand.upper() in product_brand.upper()
                            or brand.upper() in product_name.upper()
                        ):
                            brand_products.append(product)
                            if not crawl_date:
                                crawl_date = product.get("snapshot_date")

                if brand_products and crawl_date and start_date <= crawl_date <= end_date:
                    # 중복 제거 확인
                    if not any(h["date"] == crawl_date for h in sos_history):
                        # 카테고리별 총 제품 수 (Top 100 기준)
                        total_products = sum(
                            len(cat.get("products", []))
                            for cat in crawl_data.get("categories", {}).values()
                        )

                        sos = round(len(brand_products) / max(total_products, 100) * 100, 2)
                        avg_rank = round(
                            sum(p.get("rank", 0) for p in brand_products) / len(brand_products), 1
                        )

                        sos_history.append(
                            {
                                "date": crawl_date,
                                "sos": sos,
                                "product_count": len(brand_products),
                                "top10_count": sum(
                                    1 for p in brand_products if p.get("rank", 100) <= 10
                                ),
                            }
                        )
                        raw_data.append(
                            {
                                "date": crawl_date,
                                "rank": avg_rank,
                                "best_rank": min(p.get("rank", 100) for p in brand_products),
                                "worst_rank": max(p.get("rank", 100) for p in brand_products),
                            }
                        )

            except (json.JSONDecodeError, ValueError) as e:
                logging.warning(f"Failed to parse latest_crawl_result.json: {e}")

        # 3. raw_products 폴더에서 날짜별 데이터 검색 (기존 로직)
        raw_data_dir = Path("./data/raw_products")
        if raw_data_dir.exists():
            for json_file in raw_data_dir.glob("*.json"):
                try:
                    file_date = json_file.stem  # 파일명이 YYYY-MM-DD 형식이라고 가정
                    if start_date <= file_date <= end_date:
                        with open(json_file, encoding="utf-8") as f:
                            daily_raw = json.load(f)

                        # 브랜드 제품만 필터링
                        brand_products = [
                            p
                            for p in daily_raw
                            if brand.upper() in p.get("brand", "").upper()
                            or brand.upper() in p.get("product_name", "").upper()
                        ]

                        if brand_products:
                            sos = round(len(brand_products) / 100 * 100, 1)
                            avg_rank = round(
                                sum(p.get("rank", 0) for p in brand_products) / len(brand_products),
                                1,
                            )

                            # 중복 제거
                            if not any(h["date"] == file_date for h in sos_history):
                                sos_history.append(
                                    {
                                        "date": file_date,
                                        "sos": sos,
                                        "product_count": len(brand_products),
                                        "top10_count": sum(
                                            1 for p in brand_products if p.get("rank", 100) <= 10
                                        ),
                                    }
                                )
                                raw_data.append(
                                    {
                                        "date": file_date,
                                        "rank": avg_rank,
                                        "best_rank": min(
                                            p.get("rank", 100) for p in brand_products
                                        ),
                                        "worst_rank": max(
                                            p.get("rank", 100) for p in brand_products
                                        ),
                                    }
                                )
                except (json.JSONDecodeError, ValueError):
                    continue

        # 날짜순 정렬
        sos_history.sort(key=lambda x: x["date"])
        raw_data.sort(key=lambda x: x["date"])

        # available_dates 계산
        available_dates = [h["date"] for h in sos_history]

        # brand_metrics 계산 (현재 대시보드 데이터에서)
        brand_metrics = _get_brand_metrics_from_dashboard(data, brand)

        # rank_history 생성 (CPI 차트용 - 모든 브랜드 제품 포함)
        rank_history = {}
        latest_crawl_path = Path("./data/latest_crawl_result.json")
        if latest_crawl_path.exists():
            try:
                with open(latest_crawl_path, encoding="utf-8") as f:
                    crawl_data = json.load(f)
                for _cat_id, cat_data in crawl_data.get("categories", {}).items():
                    for product in cat_data.get("products", []):
                        snap_date = product.get("snapshot_date", "")
                        if not snap_date or snap_date < start_date or snap_date > end_date:
                            continue
                        if snap_date not in rank_history:
                            rank_history[snap_date] = {"products": []}
                        price_val = product.get("price", 0)
                        try:
                            price = (
                                float(str(price_val).replace("$", "").replace(",", ""))
                                if price_val
                                else 0
                            )
                        except (ValueError, TypeError):
                            price = 0
                        rank_history[snap_date]["products"].append(
                            {
                                "name": product.get("product_name", ""),
                                "brand": product.get("brand", ""),
                                "rank": product.get("rank", 0),
                                "price": price,
                            }
                        )
            except (json.JSONDecodeError, ValueError) as e:
                logging.warning(f"Failed to build rank_history from local: {e}")

        if not sos_history:
            return {
                "success": False,
                "error": "No historical data found for the specified period",
                "available_dates": [],
                "brand_metrics": [],
                "rank_history": rank_history,
                "data": None,
            }

        return {
            "success": True,
            "available_dates": available_dates,
            "brand_metrics": brand_metrics,
            "rank_history": rank_history,
            "data": {
                "sos_history": sos_history,
                "raw_data": raw_data,
                "period": {"start": start_date, "end": end_date},
                "brand": brand,
                "source": "local",
            },
        }

    except Exception as e:
        logging.error(f"Local historical data error: {e}")
        return {
            "success": False,
            "error": str(e),
            "available_dates": [],
            "brand_metrics": [],
            "data": None,
        }


@app.post("/api/export/docx")
async def export_docx(request: ExportRequest):
    """
    인사이트 리포트 DOCX 생성 및 다운로드
    """
    data = load_dashboard_data()
    if not data:
        raise HTTPException(status_code=404, detail="Dashboard data not found")

    # DOCX 문서 생성
    doc = Document()

    # 스타일 설정
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    # ===== 표지 =====
    title = doc.add_heading("AMORE INSIGHT Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("LANEIGE Amazon US 분석 리포트")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 날짜
    metadata = data.get("metadata", {})
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(
        f"분석 기준일: {metadata.get('data_date', datetime.now().strftime('%Y-%m-%d'))}"
    )
    date_para.add_run(f"\n생성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_page_break()

    # ===== 1. Executive Summary =====
    doc.add_heading("1. Executive Summary", level=1)

    brand_kpis = data.get("brand", {}).get("kpis", {})
    home_status = data.get("home", {}).get("status", {})

    summary_text = f"""
LANEIGE 브랜드는 Amazon US 시장에서 {home_status.get('exposure', 'N/A')} 상태입니다.

• Share of Shelf (SoS): {brand_kpis.get('sos', 0)}%
• Top 10 진입 제품: {brand_kpis.get('top10_count', 0)}개
• 평균 순위: {brand_kpis.get('avg_rank', 0)}위
• 시장 집중도 (HHI): {brand_kpis.get('hhi', 0)}

현재 시장 포지션: {home_status.get('position', 'N/A')}
주의 필요 제품: {home_status.get('warning_count', 0)}개
"""
    doc.add_paragraph(summary_text)

    # ===== 2. 제품별 현황 =====
    doc.add_heading("2. LANEIGE 제품 현황", level=1)

    products = data.get("products", {})
    if products:
        # 테이블 생성
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 헤더
        header_cells = table.rows[0].cells
        headers = ["제품명", "순위", "변동", "평점", "변동성"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # 데이터 행
        for _asin, product in products.items():
            row = table.add_row().cells
            row[0].text = product.get("name", "")[:40]
            row[1].text = f"#{product.get('rank', 'N/A')}"
            row[2].text = product.get("rank_delta", "-")
            row[3].text = str(product.get("rating", "-"))
            row[4].text = product.get("volatility_status", "-")

    doc.add_paragraph()

    # ===== 3. 경쟁사 분석 =====
    doc.add_heading("3. 경쟁사 분석", level=1)

    competitors = data.get("brand", {}).get("competitors", [])
    if competitors:
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"

        header_cells = table.rows[0].cells
        headers = ["브랜드", "SoS (%)", "평균 순위", "제품 수"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        for comp in competitors[:10]:
            row = table.add_row().cells
            row[0].text = comp.get("brand", "")
            row[1].text = str(comp.get("sos", 0))
            row[2].text = str(comp.get("avg_rank", "-"))
            row[3].text = str(comp.get("product_count", 0))

    doc.add_paragraph()

    # ===== 4. 액션 아이템 =====
    doc.add_heading("4. 액션 아이템", level=1)

    action_items = data.get("home", {}).get("action_items", [])
    if action_items:
        for item in action_items:
            priority_marker = "🔴" if item.get("priority") == "P1" else "🟠"
            para = doc.add_paragraph()
            para.add_run(f"{priority_marker} [{item.get('priority')}] ").bold = True
            para.add_run(f"{item.get('product_name', '')}\n")
            para.add_run(f"   신호: {item.get('signal', '')}\n")
            para.add_run(f"   권장 액션: {item.get('action_tag', '')}")
    else:
        doc.add_paragraph("현재 특별한 액션 아이템이 없습니다.")

    # ===== 5. 전략적 권고사항 =====
    if request.include_strategy:
        doc.add_heading("5. 전략적 권고사항", level=1)

        # ChatGPT로 전략 생성 (RAG 컨텍스트 활용)
        try:
            # RAG에서 전략 관련 컨텍스트 검색
            strategy_context, _ = await get_rag_context("전략 액션 권고", QueryType.ANALYSIS)

            strategy_prompt = f"""다음 데이터와 가이드라인을 바탕으로 LANEIGE 브랜드의 전략적 권고사항 3가지를 작성해주세요.

데이터:
- SoS: {brand_kpis.get('sos', 0)}%
- Top 10 제품: {brand_kpis.get('top10_count', 0)}개
- 평균 순위: {brand_kpis.get('avg_rank', 0)}위
- 주요 경쟁사: {', '.join([c['brand'] for c in competitors[:3]])}

참고 가이드라인:
{strategy_context if strategy_context else '기본 전략 기준 적용'}

각 권고사항은 1-2문장으로 간결하게 작성하세요.
"""
            response = await acompletion(
                model="gpt-4.1-mini",
                messages=[
                    {
                        "role": "system",
                        "content": "당신은 뷰티 이커머스 전문 컨설턴트입니다. 간결하고 실행 가능한 전략을 제안합니다.",
                    },
                    {"role": "user", "content": strategy_prompt},
                ],
                temperature=0.3,
                max_tokens=500,
            )

            strategy_text = response.choices[0].message.content
            doc.add_paragraph(strategy_text)

        except Exception:
            # 폴백 전략
            doc.add_paragraph("""
1. Top 10 유지 전략: 현재 상위권 제품의 리뷰 관리 및 재고 확보를 통한 포지션 유지

2. 경쟁사 모니터링: e.l.f., Maybelline 등 주요 경쟁사의 가격 및 프로모션 동향 파악

3. 신규 진입 기회: Lip Care 카테고리 외 Face Powder, Toner 등 확장 가능성 검토
""")

    # ===== 푸터 =====
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("© 2025 AMORE Pacific - Confidential").italic = True

    # BytesIO로 저장
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 파일명 생성
    filename = f"AMORE_Insight_Report_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"},
    )


@app.post("/api/export/excel")
async def export_excel(request: Request):
    """
    엑셀 데이터 내보내기 (JSON 파일 → Excel)

    데이터 소스:
    - Railway: /data/latest_crawl_result.json (Volume)
    - Local: ./data/latest_crawl_result.json
    """

    import pandas as pd

    try:
        # Parse request body
        body = await request.json()
        start_date = body.get("start_date")
        end_date = body.get("end_date")
        _include_metrics = body.get("include_metrics", True)  # reserved for future use

        # 데이터 디렉토리 경로 설정
        data_dir = Path("./data")

        # ========================================
        # 1차: SQLite에서 기간별 데이터 조회 (가장 빠름)
        # ========================================
        all_records = []
        data_source = None

        if start_date and end_date:
            # 1-1. SQLite 시도
            try:
                from src.tools.storage.sqlite_storage import get_sqlite_storage

                sqlite = get_sqlite_storage()
                await sqlite.initialize()

                # limit을 크게 설정 (5개 카테고리 × 100개 × 기간일수)
                start_dt = datetime.strptime(start_date, "%Y-%m-%d")
                end_dt = datetime.strptime(end_date, "%Y-%m-%d")
                days = (end_dt - start_dt).days + 1
                max_records = 500 * days  # 충분한 여유

                records = await sqlite.get_raw_data(
                    start_date=start_date, end_date=end_date, limit=max_records
                )

                if records:
                    all_records = records
                    data_source = "sqlite"
                    logging.info(
                        f"Excel export: loaded {len(all_records)} records from SQLite ({start_date} ~ {end_date})"
                    )

            except Exception as sqlite_err:
                logging.warning(f"Excel export: SQLite 조회 실패: {sqlite_err}")

            # 1-2. SQLite 실패 시 Google Sheets 시도
            if not all_records:
                try:
                    sheets_writer = get_sheets_writer()
                    if not sheets_writer._initialized:
                        await sheets_writer.initialize()

                    records = await sheets_writer.get_raw_data(days=days)

                    if records:
                        for record in records:
                            snapshot_date = record.get("snapshot_date", "")
                            if snapshot_date and start_date <= snapshot_date <= end_date:
                                all_records.append(record)

                        if all_records:
                            data_source = "sheets"
                            logging.info(
                                f"Excel export: loaded {len(all_records)} records from Google Sheets ({start_date} ~ {end_date})"
                            )

                except Exception as sheets_err:
                    logging.warning(f"Excel export: Google Sheets 조회 실패: {sheets_err}")

        # ========================================
        # 2차: 로컬 JSON 파일에서 데이터 로드 (폴백)
        # ========================================
        crawl_data = None
        json_path = None

        if not data_source:
            possible_paths = [
                data_dir / "latest_crawl_result.json",
                data_dir / "dashboard_data.json",
            ]

            for path in possible_paths:
                if path.exists():
                    json_path = path
                    break

            if json_path is None:
                raise HTTPException(
                    status_code=404, detail="크롤링 데이터가 없습니다. 먼저 크롤링을 실행해주세요."
                )

            with open(json_path, encoding="utf-8") as f:
                crawl_data = json.load(f)

            logging.info(f"Excel export: loaded data from {json_path}")

        # 데이터 소스 유형 판단
        # 1. data_source: SQLite 또는 Google Sheets에서 기간별 데이터 로드됨
        # 2. is_crawl_data: latest_crawl_result.json의 raw 데이터
        # 3. is_dashboard_data: dashboard_data.json의 집계 데이터
        is_crawl_data = False
        is_dashboard_data = False

        if crawl_data:
            if "categories" in crawl_data:
                first_cat = next(iter(crawl_data["categories"].values()), {})
                is_crawl_data = isinstance(first_cat, dict) and (
                    "rank_records" in first_cat or "products" in first_cat
                )
            is_dashboard_data = "metadata" in crawl_data and "brand" in crawl_data

        # 출력 경로 (Railway 환경 고려)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_dir = data_dir / "exports"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"AMORE_Data_{timestamp}.xlsx"

        sheets_created = []
        total_rows = 0

        # 카테고리 매핑
        categories_info = {
            "beauty": "Beauty & Personal Care",
            "skin_care": "Skin Care",
            "lip_care": "Lip Care",
            "lip_makeup": "Lip Makeup",
            "face_powder": "Face Powder",
        }

        with pd.ExcelWriter(str(output_path), engine="openpyxl") as writer:
            # Google Sheets RawData와 동일한 컬럼 순서
            RAWDATA_COLUMNS = [
                "snapshot_date",
                "category_id",
                "rank",
                "asin",
                "product_name",
                "brand",
                "price",
                "list_price",
                "discount_percent",
                "rating",
                "reviews_count",
                "badge",
                "coupon_text",
                "is_subscribe_save",
                "promo_badges",
                "product_url",
            ]

            # ========================================
            # Case 1: SQLite/Google Sheets에서 기간별 데이터 로드됨
            # ========================================
            if data_source and all_records:
                source_name = "SQLite" if data_source == "sqlite" else "Google Sheets"
                logging.info(
                    f"Excel export: using {source_name} data ({len(all_records)} records, {start_date} ~ {end_date})"
                )

                df_all = pd.DataFrame(all_records)

                if not df_all.empty:
                    # 1. RawData 시트 - 전체 데이터
                    available_cols = [c for c in RAWDATA_COLUMNS if c in df_all.columns]
                    df_raw = df_all[available_cols].copy()
                    df_raw = df_raw.sort_values(["snapshot_date", "category_id", "rank"])
                    df_raw.to_excel(writer, sheet_name="RawData", index=False)
                    sheets_created.append("RawData")
                    total_rows += len(df_raw)

                    # 2. 날짜별 요약 시트
                    if "snapshot_date" in df_all.columns:
                        date_summary = []
                        for date in sorted(df_all["snapshot_date"].unique()):
                            df_date = df_all[df_all["snapshot_date"] == date]
                            laneige_count = (
                                len(df_date[df_date["brand"].str.upper() == "LANEIGE"])
                                if "brand" in df_date.columns
                                else 0
                            )
                            date_summary.append(
                                {
                                    "날짜": date,
                                    "총 제품 수": len(df_date),
                                    "LANEIGE 제품 수": laneige_count,
                                    "LANEIGE SoS (%)": round(laneige_count / len(df_date) * 100, 1)
                                    if len(df_date) > 0
                                    else 0,
                                }
                            )
                        if date_summary:
                            df_summary = pd.DataFrame(date_summary)
                            df_summary.to_excel(writer, sheet_name="Daily Summary", index=False)
                            sheets_created.append("Daily Summary")
                            total_rows += len(df_summary)

                    # 3. 카테고리별 시트
                    if "category_id" in df_all.columns:
                        for cat_id in df_all["category_id"].unique():
                            df_cat = df_all[df_all["category_id"] == cat_id].copy()
                            if df_cat.empty:
                                continue

                            display_cols = [
                                "snapshot_date",
                                "rank",
                                "asin",
                                "product_name",
                                "brand",
                                "price",
                                "rating",
                                "reviews_count",
                                "badge",
                            ]
                            available_display = [c for c in display_cols if c in df_cat.columns]
                            df_display = df_cat[available_display].sort_values(
                                ["snapshot_date", "rank"]
                            )

                            sheet_name = categories_info.get(cat_id, cat_id)[:31]
                            df_display.to_excel(writer, sheet_name=sheet_name, index=False)
                            sheets_created.append(sheet_name)
                            total_rows += len(df_display)

                    # 4. LANEIGE 제품 전용 시트
                    if "brand" in df_all.columns:
                        df_laneige = df_all[df_all["brand"].str.upper() == "LANEIGE"].copy()
                        if not df_laneige.empty:
                            laneige_cols = [
                                "snapshot_date",
                                "category_id",
                                "rank",
                                "asin",
                                "product_name",
                                "price",
                                "rating",
                                "reviews_count",
                                "badge",
                            ]
                            available_laneige = [c for c in laneige_cols if c in df_laneige.columns]
                            df_laneige = df_laneige[available_laneige].sort_values(
                                ["snapshot_date", "category_id", "rank"]
                            )
                            df_laneige.to_excel(writer, sheet_name="LANEIGE Products", index=False)
                            sheets_created.append("LANEIGE Products")
                            total_rows += len(df_laneige)

            # ========================================
            # Case 2: 대시보드 데이터 형식 (집계 데이터만)
            # ========================================
            elif is_dashboard_data and not is_crawl_data:
                logging.info("Excel export: using dashboard_data.json (aggregated data only)")

                # 1. Overview 시트
                metadata = crawl_data.get("metadata", {})
                data_source = crawl_data.get("data_source", {})
                overview_data = [
                    {"항목": "데이터 날짜", "값": metadata.get("data_date", "N/A")},
                    {"항목": "생성 시각", "값": metadata.get("generated_at", "N/A")},
                    {"항목": "총 제품 수", "값": metadata.get("total_products", 0)},
                    {"항목": "LANEIGE 제품 수", "값": metadata.get("laneige_products", 0)},
                    {"항목": "플랫폼", "값": data_source.get("platform", "Amazon US")},
                ]
                df_overview = pd.DataFrame(overview_data)
                df_overview.to_excel(writer, sheet_name="Overview", index=False)
                sheets_created.append("Overview")
                total_rows += len(df_overview)

                # 2. Brand KPIs 시트
                brand_kpis = crawl_data.get("brand", {}).get("kpis", {})
                if brand_kpis:
                    kpi_data = [
                        {"KPI": "SoS (Share of Shelf)", "값": f"{brand_kpis.get('sos', 0)}%"},
                        {"KPI": "SoS 변화", "값": brand_kpis.get("sos_delta", "N/A")},
                        {"KPI": "Top 10 제품 수", "값": brand_kpis.get("top10_count", 0)},
                        {"KPI": "평균 순위", "값": brand_kpis.get("avg_rank", 0)},
                        {"KPI": "HHI (시장 집중도)", "값": brand_kpis.get("hhi", 0)},
                    ]
                    df_kpis = pd.DataFrame(kpi_data)
                    df_kpis.to_excel(writer, sheet_name="LANEIGE KPIs", index=False)
                    sheets_created.append("LANEIGE KPIs")
                    total_rows += len(df_kpis)

                # 3. Competitors 시트
                competitors = crawl_data.get("brand", {}).get("competitors", [])
                if competitors:
                    df_comp = pd.DataFrame(competitors)
                    column_mapping = {
                        "brand": "Brand",
                        "sos": "SoS (%)",
                        "avg_rank": "Avg Rank",
                        "product_count": "Product Count",
                        "avg_price": "Avg Price ($)",
                    }
                    existing_cols = {
                        k: v for k, v in column_mapping.items() if k in df_comp.columns
                    }
                    df_comp = df_comp.rename(columns=existing_cols)
                    df_comp.to_excel(writer, sheet_name="Competitors", index=False)
                    sheets_created.append("Competitors")
                    total_rows += len(df_comp)

                # 4. Action Items 시트
                action_items = crawl_data.get("home", {}).get("action_items", [])
                if action_items:
                    df_actions = pd.DataFrame(action_items)
                    df_actions.to_excel(writer, sheet_name="Action Items", index=False)
                    sheets_created.append("Action Items")
                    total_rows += len(df_actions)

                # 5. Category View 시트
                category_data = crawl_data.get("category", {})
                if category_data:
                    for cat_id, cat_info in category_data.items():
                        top_products = cat_info.get("top_products", [])
                        if top_products:
                            df_cat = pd.DataFrame(top_products)
                            sheet_name = categories_info.get(cat_id, cat_id)[:31]
                            df_cat.to_excel(writer, sheet_name=sheet_name, index=False)
                            sheets_created.append(sheet_name)
                            total_rows += len(df_cat)

            # ========================================
            # Case 3: 로컬 크롤링 원본 데이터
            # ========================================
            else:
                logging.info("Excel export: using latest_crawl_result.json (raw crawl data)")

                # 전체 RawData 수집 (카테고리별 rank_records)
                all_records = []
                for cat_id, cat_data in crawl_data.get("categories", {}).items():
                    records = cat_data.get("rank_records", cat_data.get("products", []))
                    for record in records:
                        # category_id 추가 (없는 경우)
                        if "category_id" not in record:
                            record["category_id"] = cat_id
                        all_records.append(record)

                if not all_records:
                    logging.warning("Excel export: no rank_records found in crawl data")

                if all_records:
                    df_all = pd.DataFrame(all_records)

                    # 날짜 필터 적용 (선택 기간)
                    if start_date and "snapshot_date" in df_all.columns:
                        df_all = df_all[df_all["snapshot_date"] >= start_date]
                    if end_date and "snapshot_date" in df_all.columns:
                        df_all = df_all[df_all["snapshot_date"] <= end_date]

                    if not df_all.empty:
                        # 1. RawData 시트 - Google Sheets와 동일한 전체 데이터
                        available_cols = [c for c in RAWDATA_COLUMNS if c in df_all.columns]
                        df_raw = df_all[available_cols].copy()
                        df_raw = df_raw.sort_values(["category_id", "rank"])
                        df_raw.to_excel(writer, sheet_name="RawData", index=False)
                        sheets_created.append("RawData")
                        total_rows += len(df_raw)

                        # 2. 카테고리별 시트 (요약 보기용)
                        for cat_id in df_all["category_id"].unique():
                            df_cat = df_all[df_all["category_id"] == cat_id].copy()
                            if df_cat.empty:
                                continue

                            # 핵심 컬럼만 선택하여 가독성 향상
                            display_cols = [
                                "rank",
                                "asin",
                                "product_name",
                                "brand",
                                "price",
                                "rating",
                                "reviews_count",
                                "badge",
                            ]
                            available_display = [c for c in display_cols if c in df_cat.columns]
                            df_display = df_cat[available_display].sort_values("rank")

                            # 시트 이름 (31자 제한)
                            sheet_name = categories_info.get(cat_id, cat_id)[:31]
                            df_display.to_excel(writer, sheet_name=sheet_name, index=False)
                            sheets_created.append(sheet_name)
                            total_rows += len(df_display)

                        # 3. LANEIGE 제품 전용 시트
                        df_laneige = df_all[df_all["brand"].str.upper() == "LANEIGE"].copy()
                        if not df_laneige.empty:
                            laneige_cols = [
                                "snapshot_date",
                                "category_id",
                                "rank",
                                "asin",
                                "product_name",
                                "price",
                                "rating",
                                "reviews_count",
                                "badge",
                            ]
                            available_laneige = [c for c in laneige_cols if c in df_laneige.columns]
                            df_laneige = df_laneige[available_laneige].sort_values(
                                ["category_id", "rank"]
                            )
                            df_laneige.to_excel(writer, sheet_name="LANEIGE Products", index=False)
                            sheets_created.append("LANEIGE Products")
                            total_rows += len(df_laneige)

                        # 4. Summary 시트 - 브랜드별 집계
                        if "brand" in df_all.columns:
                            agg_dict = {"asin": "count"}
                            if "rank" in df_all.columns:
                                agg_dict["rank"] = "mean"
                            if "price" in df_all.columns:
                                agg_dict["price"] = "mean"
                            if "rating" in df_all.columns:
                                agg_dict["rating"] = "mean"

                            summary = df_all.groupby("brand").agg(agg_dict).reset_index()
                            col_names = ["Brand", "Product Count"]
                            if "rank" in agg_dict:
                                col_names.append("Avg Rank")
                            if "price" in agg_dict:
                                col_names.append("Avg Price")
                            if "rating" in agg_dict:
                                col_names.append("Avg Rating")
                            summary.columns = col_names

                            summary = summary.sort_values("Product Count", ascending=False).head(30)
                            for col in ["Avg Rank", "Avg Price", "Avg Rating"]:
                                if col in summary.columns:
                                    summary[col] = summary[col].round(2)

                            summary.to_excel(writer, sheet_name="Summary", index=False)
                            sheets_created.append("Summary")
                            total_rows += len(summary)

            # 4. 시트가 하나도 없으면 안내 시트 생성
            if not sheets_created:
                data_source_info = (
                    "SQLite"
                    if data_source == "sqlite"
                    else (
                        "Google Sheets"
                        if data_source == "sheets"
                        else (str(json_path) if json_path else "N/A")
                    )
                )
                no_data_info = [
                    {"항목": "요청 기간", "값": f"{start_date or 'N/A'} ~ {end_date or 'N/A'}"},
                    {"항목": "결과", "값": "해당 기간에 데이터가 없습니다"},
                    {"항목": "데이터 소스", "값": data_source_info},
                    {"항목": "안내", "값": "크롤링 실행 후 다시 시도해주세요"},
                ]
                df_no_data = pd.DataFrame(no_data_info)
                df_no_data.to_excel(writer, sheet_name="No Data", index=False)
                sheets_created.append("No Data")

        logging.info(f"Excel exported: {output_path} ({total_rows} rows, sheets: {sheets_created})")

        # 파일 반환
        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={output_path.name}"},
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Excel export error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Excel 내보내기 중 오류가 발생했습니다") from e


# ============= Alert Settings API =============

from src.api.dependencies import get_app_state_manager
from src.core.state_manager import EmailSubscription


@app.get("/api/v3/alert-settings")
async def get_alert_settings():
    """
    현재 알림 설정 조회

    참고: 현재는 단일 사용자 설정만 지원 (첫 번째 등록된 이메일)
    """
    state_manager = get_app_state_manager()
    subscriptions = state_manager.get_all_subscriptions()

    if not subscriptions:
        return {"email": "", "consent": False, "alert_types": [], "consent_date": None}

    # 첫 번째 구독 반환
    email, sub = next(iter(subscriptions.items()))
    return {
        "email": email,
        "consent": sub.consent,
        "alert_types": sub.alert_types,
        "consent_date": sub.consent_date.isoformat() if sub.consent_date else None,
    }


@app.post("/api/v3/alert-settings", dependencies=[Depends(verify_api_key)])
@limiter.limit("5/minute")  # 분당 5회 제한 (스팸 방지)
async def save_alert_settings(request: Request, settings: AlertSettingsRequest):
    """
    알림 설정 저장

    보안: API Key + Rate Limiting (IP당 분당 5회)
    중요: consent가 True일 때만 이메일 등록
    """
    state_manager = get_app_state_manager()

    if not settings.email:
        raise HTTPException(status_code=400, detail="이메일 주소가 필요합니다.")

    if settings.consent:
        # 이메일 등록 (명시적 동의)
        success = state_manager.register_email(
            email=settings.email, consent=True, alert_types=settings.alert_types
        )

        if not success:
            raise HTTPException(status_code=400, detail="이메일 등록 실패")

        return {"status": "ok", "message": "알림 설정이 저장되었습니다."}
    else:
        # 동의 없으면 업데이트만 (알림 유형 변경)
        success = state_manager.update_email_subscription(
            email=settings.email, alert_types=settings.alert_types
        )

        return {"status": "ok", "message": "설정이 업데이트되었습니다."}


@app.post("/api/v3/alert-settings/revoke", dependencies=[Depends(verify_api_key)])
@limiter.limit("5/minute")  # 분당 5회 제한
async def revoke_alert_consent(request: Request):
    """
    알림 동의 철회

    보안: API Key + Rate Limiting
    첫 번째 등록된 이메일의 동의를 철회합니다.
    """
    state_manager = get_app_state_manager()
    subscriptions = state_manager.get_all_subscriptions()

    if not subscriptions:
        return {"status": "ok", "message": "철회할 동의가 없습니다."}

    # 첫 번째 이메일 철회
    email = next(iter(subscriptions.keys()))
    state_manager.revoke_email_consent(email)

    return {"status": "ok", "message": "동의가 철회되었습니다."}


# =============================================================================
# v4 Alert Settings API (뉴닉 스타일 구독 플로우)
# =============================================================================


@app.post("/api/v4/subscribe")
@limiter.limit("3/minute")
async def subscribe_v4(request: Request, body: SubscribeRequest):
    """
    구독 시작 (v4 통합 엔드포인트)

    - 신규 이메일: JWT 인증 메일 발송 + alert_types 임시 저장
    - 기존 이메일 (already_verified): 현재 구독 설정 반환
    """
    import re

    email = body.email.strip()
    email_regex = r"^[^\s@]+@[^\s@]+\.[^\s@]+$"
    if not email or not re.match(email_regex, email):
        raise HTTPException(status_code=400, detail="올바른 이메일 주소를 입력해주세요.")

    if not body.alert_types:
        raise HTTPException(status_code=400, detail="최소 하나 이상의 알림 유형을 선택해주세요.")

    state_manager = get_state_manager()
    existing = state_manager.get_subscription(email)

    # 이미 인증된 이메일
    if existing and existing.verified:
        return {
            "success": True,
            "already_verified": True,
            "message": "이미 가입한 이메일이에요.",
            "current_settings": {
                "alert_types": existing.alert_types,
                "active": existing.active,
                "consent": existing.consent,
            },
        }

    # 신규 이메일 - JWT 인증 메일 발송
    try:
        token = create_email_verification_token(email)

        base_url = get_base_url()
        verify_url = f"{base_url}/api/alerts/confirm-email?token={token}&email={email}"

        from src.tools.notifications.email_sender import EmailSender

        email_sender = EmailSender()

        if not email_sender.is_enabled():
            raise HTTPException(status_code=503, detail="이메일 서비스가 설정되지 않았습니다.")

        result = await email_sender.send_verification_email(
            recipient=email, verify_url=verify_url, token=token
        )

        if result.success:
            # 인증 전이지만 선택한 alert_types를 미리 저장 (인증 완료 시 적용)
            if not existing:
                # 새 구독 생성 (아직 미인증, 미동의 상태)
                sub = EmailSubscription(
                    email=email,
                    consent=False,
                    alert_types=body.alert_types,
                    active=False,
                    verified=False,
                )
                state_manager._email_subscriptions[email] = sub
                state_manager._save_subscriptions()
            else:
                # 기존 미인증 구독 업데이트
                existing.alert_types = body.alert_types
                state_manager._save_subscriptions()

            logging.info(f"[v4] Verification email sent to {email}, alert_types={body.alert_types}")
            return {
                "success": True,
                "already_verified": False,
                "message": "인증 이메일이 발송되었습니다. (30분 내 인증해주세요)",
            }
        else:
            raise HTTPException(status_code=500, detail=f"이메일 발송 실패: {result.message}")

    except ValueError as e:
        logging.error(f"JWT configuration error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"[v4] Subscribe error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@app.get("/api/v4/alert-settings")
async def get_alert_settings_v4(email: str | None = None):
    """
    알림 설정 조회 (v4)

    Args:
        email: 조회할 이메일 (없으면 첫 번째 구독자)
    """
    state_manager = get_state_manager()

    if email:
        sub = state_manager.get_subscription(email)
        if not sub:
            return {"found": False, "email": email, "message": "등록되지 않은 이메일입니다."}
        return {
            "found": True,
            "email": sub.email,
            "consent": sub.consent,
            "alert_types": sub.alert_types,
            "active": sub.active,
            "verified": sub.verified,
            "verified_at": sub.verified_at.isoformat() if sub.verified_at else None,
            "consent_date": sub.consent_date.isoformat() if sub.consent_date else None,
        }

    # email 미지정 시 기존 v3 동작 (첫 번째 구독자)
    subscriptions = state_manager.get_all_subscriptions()
    if not subscriptions:
        return {"found": False, "email": "", "consent": False, "alert_types": []}

    email_key, sub = next(iter(subscriptions.items()))
    return {
        "found": True,
        "email": email_key,
        "consent": sub.consent,
        "alert_types": sub.alert_types,
        "active": sub.active,
        "verified": sub.verified,
        "verified_at": sub.verified_at.isoformat() if sub.verified_at else None,
        "consent_date": sub.consent_date.isoformat() if sub.consent_date else None,
    }


@app.put("/api/v4/alert-settings")
@limiter.limit("5/minute")
async def update_alert_settings_v4(request: Request, body: UpdateAlertSettingsRequest):
    """
    알림 설정 수정 (v4) - 기존 구독자 전용

    인증 완료된 이메일만 수정 가능
    """
    email = body.email.strip()
    if not email:
        raise HTTPException(status_code=400, detail="이메일 주소가 필요합니다.")

    state_manager = get_state_manager()
    sub = state_manager.get_subscription(email)

    if not sub:
        raise HTTPException(status_code=404, detail="등록되지 않은 이메일입니다.")

    if not sub.verified:
        raise HTTPException(status_code=403, detail="이메일 인증이 완료되지 않았습니다.")

    # alert_types 업데이트
    success = state_manager.update_email_subscription(
        email=email, alert_types=body.alert_types, active=True
    )

    # consent도 True로 설정 (설정 수정 = 동의 유지)
    if success and not sub.consent:
        sub.consent = True
        sub.consent_date = datetime.now()
        state_manager._save_subscriptions()

    if success:
        return {
            "status": "ok",
            "message": "알림 설정이 업데이트되었습니다.",
            "alert_types": body.alert_types,
        }
    else:
        raise HTTPException(status_code=500, detail="설정 업데이트 실패")


@app.delete("/api/v4/alert-settings")
@limiter.limit("5/minute")
async def delete_alert_settings_v4(request: Request, email: str):
    """
    구독 해지 (v4)

    Args:
        email: 해지할 이메일 주소
    """
    if not email:
        raise HTTPException(status_code=400, detail="이메일 주소가 필요합니다.")

    state_manager = get_state_manager()
    sub = state_manager.get_subscription(email)

    if not sub:
        raise HTTPException(status_code=404, detail="등록되지 않은 이메일입니다.")

    state_manager.revoke_email_consent(email)
    return {"status": "ok", "message": "구독이 해지되었습니다."}


@app.get("/api/v3/alerts")
async def get_alerts(limit: int = 50, alert_type: str | None = None):
    """
    알림 목록 조회

    Args:
        limit: 최대 개수
        alert_type: 필터할 알림 유형
    """
    from src.agents.alert_agent import AlertAgent

    state_manager = get_app_state_manager()
    alert_agent = AlertAgent(state_manager)

    return {
        "alerts": alert_agent.get_alerts(limit=limit, alert_type=alert_type),
        "pending_count": alert_agent.get_pending_count(),
        "stats": alert_agent.get_stats(),
    }


# ============= Level 4 Brain API (v4) =============


@app.post("/api/v4/chat", response_model=BrainChatResponse, dependencies=[Depends(verify_api_key)])
@limiter.limit("10/minute")  # 분당 10회 제한 (보안 강화)
async def chat_v4(request: Request, body: BrainChatRequest):
    """
    Level 4 Brain 기반 챗봇 API (v4)

    LLM-First 접근:
    - 모든 판단을 LLM이 수행
    - 규칙 기반 빠른 경로 없음
    - RAG + KG 하이브리드 검색
    - 자율 스케줄러와 통합
    """
    import time

    start_time = time.time()

    message = body.message.strip()
    session_id = body.session_id or "default"

    if not message:
        raise HTTPException(status_code=400, detail="Message is required")

    try:
        # Brain 인스턴스 획득
        brain = await get_initialized_brain()

        # 현재 메트릭 데이터 로드
        data = load_dashboard_data()
        current_metrics = data if data else None

        # Brain으로 처리 (LLM-First)
        response = await brain.process_query(
            query=message,
            session_id=session_id,
            current_metrics=current_metrics,
            skip_cache=body.skip_cache,
        )

        processing_time = (time.time() - start_time) * 1000

        return BrainChatResponse(
            text=response.text,
            confidence=response.confidence_score,
            sources=response.sources if isinstance(response.sources, list) else [],
            reasoning=response.query_type,
            tools_used=response.tools_called,
            processing_time_ms=processing_time,
            from_cache=False,
            brain_mode=brain.mode.value,
            suggestions=response.suggestions,
            query_type=response.query_type,
        )

    except Exception as e:
        logging.error(f"Brain error: {e}")
        return BrainChatResponse(
            text=f"처리 중 오류가 발생했습니다: {str(e)}",
            confidence=0.0,
            sources=[],
            reasoning=None,
            tools_used=[],
            processing_time_ms=(time.time() - start_time) * 1000,
            from_cache=False,
            brain_mode="error",
        )


@app.post("/api/v4/chat/stream")
@limiter.limit("10/minute")
async def chat_v4_stream(request: Request, body: BrainChatRequest):
    """
    Level 4 Brain 기반 SSE 스트리밍 챗봇 API (v4)

    v3의 SSE 스트리밍과 동일한 인터페이스로 v4 Brain의 처리 결과를 반환합니다.
    ReAct + OWL + PromptGuard + 도구 호출을 모두 지원합니다.

    이벤트 타입:
    - status: 처리 단계 알림
    - tool_call: 도구 호출 정보
    - text: 응답 텍스트
    - done: 완료 (메타데이터 포함)
    - error: 오류 발생
    """
    message = body.message.strip()
    session_id = body.session_id or "default"

    if not message:
        raise HTTPException(status_code=400, detail="Message is required")

    try:
        brain = await get_initialized_brain()
        data = load_dashboard_data()
        current_metrics = data if data else None

        async def generate():
            try:
                async for chunk in brain.process_query_stream(
                    query=message,
                    session_id=session_id,
                    current_metrics=current_metrics,
                ):
                    event_data = json.dumps(chunk, ensure_ascii=False)
                    yield f"data: {event_data}\n\n"
            except Exception as e:
                logger.error(f"v4 SSE stream error: {e}")
                error_data = json.dumps({"type": "error", "content": str(e)}, ensure_ascii=False)
                yield f"data: {error_data}\n\n"

        return StreamingResponse(
            generate(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )

    except Exception as e:
        logger.error(f"v4 chat stream init error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


# ============= Amazon Deals API =============

from src.tools.scrapers.deals_scraper import get_deals_scraper


@app.get("/api/deals")
async def get_deals_data(brand: str | None = None, hours: int = 24, limit: int = 100):
    """
    저장된 Deals 데이터 조회

    Args:
        brand: 브랜드 필터 (선택)
        hours: 최근 N시간 데이터 (기본: 24시간)
        limit: 최대 개수

    Returns:
        - deals: 딜 데이터 리스트
        - summary: 요약 통계
    """
    try:
        storage = get_sqlite_storage()
        await storage.initialize()

        # 경쟁사 딜 조회
        deals = await storage.get_competitor_deals(brand=brand, hours=hours)

        # 최대 개수 제한
        deals = deals[:limit] if len(deals) > limit else deals

        # 요약 통계
        summary = await storage.get_deals_summary(days=7)

        return {
            "success": True,
            "deals": deals,
            "count": len(deals),
            "summary": summary,
            "filters": {"brand": brand, "hours": hours},
        }

    except Exception as e:
        logging.error(f"Deals data error: {e}")
        return {"success": False, "deals": [], "count": 0, "error": str(e)}


@app.get("/api/deals/summary")
async def get_deals_summary(days: int = 7):
    """
    Deals 요약 통계

    Args:
        days: 분석 기간 (일)

    Returns:
        - by_brand: 브랜드별 딜 현황
        - by_date: 일별 추이
    """
    try:
        storage = get_sqlite_storage()
        await storage.initialize()

        summary = await storage.get_deals_summary(days=days)

        return {"success": True, **summary}

    except Exception as e:
        logging.error(f"Deals summary error: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/deals/scrape", dependencies=[Depends(verify_api_key)])
async def scrape_deals(request: DealsRequest):
    """
    Amazon Deals 페이지 크롤링 (API Key 필요)

    경쟁사 할인 정보를 수집하고 저장합니다.

    Args:
        max_items: 최대 수집 개수
        beauty_only: 뷰티 카테고리만 필터링

    Returns:
        - deals: 수집된 딜 데이터
        - competitor_deals: 경쟁사 딜
        - lightning_count: Lightning Deal 수
    """
    try:
        scraper = await get_deals_scraper()

        # 크롤링 실행
        result = await scraper.scrape_deals(
            max_items=request.max_items, beauty_only=request.beauty_only
        )

        if result["success"]:
            # SQLite에 저장
            storage = get_sqlite_storage()
            await storage.initialize()

            # 모든 딜 저장
            if result["deals"]:
                await storage.save_deals(result["deals"], is_competitor=False)

            # 경쟁사 딜은 is_competitor=True로 별도 저장
            if result["competitor_deals"]:
                await storage.save_deals(result["competitor_deals"], is_competitor=True)

                # 알림 서비스로 알림 처리
                try:
                    alert_service = get_alert_service()
                    alerts = await alert_service.process_deals_for_alerts(
                        result["competitor_deals"]
                    )

                    # DB에 알림 저장
                    for alert in alerts:
                        await storage.save_deal_alert(alert)

                    logging.info(
                        f"Processed {len(alerts)} alerts from {len(result['competitor_deals'])} competitor deals"
                    )
                except Exception as alert_err:
                    logging.error(f"Alert processing error: {alert_err}")
                    # 알림 실패해도 크롤링 결과는 반환

            logging.info(
                f"Deals scraped: {result['count']} total, {len(result['competitor_deals'])} competitors"
            )

        return DealsResponse(
            success=result["success"],
            count=result["count"],
            lightning_count=result["lightning_count"],
            competitor_count=len(result["competitor_deals"]),
            snapshot_datetime=result["snapshot_datetime"],
            deals=result["deals"],
            competitor_deals=result["competitor_deals"],
            error=result.get("error"),
        )

    except Exception as e:
        logging.error(f"Deals scrape error: {e}")
        return DealsResponse(
            success=False,
            count=0,
            lightning_count=0,
            competitor_count=0,
            snapshot_datetime=datetime.now().isoformat(),
            deals=[],
            competitor_deals=[],
            error=str(e),
        )


@app.get("/api/deals/alerts")
async def get_deals_alerts(limit: int = 50, unsent_only: bool = False):
    """
    Deals 알림 목록 조회

    Args:
        limit: 최대 개수
        unsent_only: 미발송 알림만 조회

    Returns:
        - alerts: 알림 목록
        - count: 총 개수
    """
    try:
        storage = get_sqlite_storage()
        await storage.initialize()

        if unsent_only:
            alerts = await storage.get_unsent_alerts(limit=limit)
        else:
            # 모든 알림 조회 (최근 7일)
            with storage.get_connection() as conn:
                cursor = conn.execute(
                    """
                    SELECT * FROM deals_alerts
                    ORDER BY alert_datetime DESC
                    LIMIT ?
                """,
                    (limit,),
                )
                alerts = [dict(row) for row in cursor.fetchall()]

        return {"success": True, "alerts": alerts, "count": len(alerts)}

    except Exception as e:
        logging.error(f"Deals alerts error: {e}")
        return {"success": False, "alerts": [], "count": 0, "error": str(e)}


@app.post("/api/deals/export")
async def export_deals_report(days: int = 7, format: str = "excel"):
    """
    Deals 리포트 내보내기

    Args:
        days: 분석 기간 (일)
        format: 출력 형식 (excel, json)

    Returns:
        - 엑셀: 파일 다운로드
        - JSON: 데이터 반환
    """
    try:
        storage = get_sqlite_storage()
        await storage.initialize()

        if format == "json":
            # JSON 형식 반환
            summary = await storage.get_deals_summary(days=days)

            # 전체 딜 데이터
            with storage.get_connection() as conn:
                cutoff_date = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%d")
                cursor = conn.execute(
                    """
                    SELECT * FROM deals
                    WHERE DATE(snapshot_datetime) >= ?
                    ORDER BY snapshot_datetime DESC
                """,
                    (cutoff_date,),
                )
                all_deals = [dict(row) for row in cursor.fetchall()]

            return {
                "success": True,
                "summary": summary,
                "deals": all_deals,
                "export_date": datetime.now().isoformat(),
                "period_days": days,
            }

        else:  # Excel
            # 엑셀 파일 생성
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"./data/exports/Deals_Report_{timestamp}.xlsx"

            result = storage.export_deals_report(output_path=output_path, days=days)

            if not result.get("success"):
                raise HTTPException(status_code=500, detail=result.get("error", "Export failed"))

            file_path = Path(result["file_path"])
            if not file_path.exists():
                raise HTTPException(status_code=500, detail="Generated file not found")

            return FileResponse(
                path=str(file_path),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={"Content-Disposition": f"attachment; filename={file_path.name}"},
            )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Deals export error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Deals 내보내기 중 오류가 발생했습니다") from e


# ============= 알림 서비스 API =============

from src.tools.notifications.alert_service import get_alert_service


@app.get("/api/alerts/status")
async def get_alert_service_status():
    """알림 서비스 상태 조회"""
    try:
        service = get_alert_service()
        return {"success": True, **service.get_status()}
    except Exception as e:
        logging.error(f"Alert service status error: {e}")
        return {"success": False, "error": str(e)}


@app.post("/api/alerts/send")
async def send_pending_alerts(request: AlertSendRequest | None = None):
    """
    미발송 알림 발송

    특정 alert_ids를 지정하면 해당 알림만, 없으면 미발송 전체 발송
    """
    try:
        storage = get_sqlite_storage()
        await storage.initialize()

        alert_service = get_alert_service()

        # 미발송 알림 조회
        unsent_alerts = await storage.get_unsent_alerts(limit=50)

        if not unsent_alerts:
            return {"success": True, "message": "No pending alerts to send", "sent_count": 0}

        # 특정 ID 필터링
        if request and request.alert_ids:
            unsent_alerts = [a for a in unsent_alerts if a.get("id") in request.alert_ids]

        if not unsent_alerts:
            return {"success": True, "message": "No matching alerts found", "sent_count": 0}

        # 알림 발송
        sent_count = 0
        for alert in unsent_alerts:
            result = await alert_service.send_single_alert(alert)

            # 성공 시 발송 완료 표시
            if result.get("slack") or result.get("email"):
                await storage.mark_alert_sent(alert["id"])
                sent_count += 1

        return {
            "success": True,
            "sent_count": sent_count,
            "total_pending": len(unsent_alerts),
            "channels": {
                "slack": alert_service._slack_enabled,
                "email": alert_service._email_enabled,
            },
        }

    except Exception as e:
        logging.error(f"Alert send error: {e}")
        return {"success": False, "error": str(e), "sent_count": 0}


@app.post("/api/alerts/test")
async def send_test_alert():
    """테스트 알림 발송"""
    try:
        alert_service = get_alert_service()

        test_alert = {
            "alert_datetime": datetime.now().isoformat(),
            "brand": "TEST BRAND",
            "asin": "B000TEST01",
            "product_name": "Test Product - Alert System Verification",
            "deal_type": "lightning",
            "discount_percent": 50.0,
            "deal_price": 19.99,
            "original_price": 39.99,
            "time_remaining": "2h 30m",
            "claimed_percent": 45,
            "product_url": "https://amazon.com/dp/B000TEST01",
            "alert_type": "lightning_deal",
            "alert_message": "Test Alert - 시스템 테스트 알림입니다",
        }

        result = await alert_service.send_single_alert(test_alert)

        return {
            "success": True,
            "test_alert": test_alert,
            "send_result": result,
            "message": "Test alert sent successfully"
            if any(result.values())
            else "No channels enabled",
        }

    except Exception as e:
        logging.error(f"Test alert error: {e}")
        return {"success": False, "error": str(e)}


# ============= Email Verification API =============

import jwt

# JWT 설정
JWT_SECRET_KEY = os.getenv("JWT_SECRET_KEY")
JWT_ALGORITHM = "HS256"
EMAIL_VERIFICATION_EXPIRES_MINUTES = 30  # 30분 만료


def create_email_verification_token(
    email: str, expires_minutes: int = EMAIL_VERIFICATION_EXPIRES_MINUTES
) -> str:
    """
    이메일 인증용 JWT 토큰 생성

    Args:
        email: 인증할 이메일 주소
        expires_minutes: 토큰 만료 시간 (분)

    Returns:
        JWT 토큰 문자열
    """
    if not JWT_SECRET_KEY:
        raise ValueError("JWT_SECRET_KEY 환경변수가 설정되지 않았습니다.")

    payload = {
        "email": email,
        "purpose": "email_verification",
        "exp": datetime.now(UTC) + timedelta(minutes=expires_minutes),
        "iat": datetime.now(UTC),
    }
    return jwt.encode(payload, JWT_SECRET_KEY, algorithm=JWT_ALGORITHM)


def verify_jwt_email_token(token: str) -> dict:
    """
    JWT 이메일 인증 토큰 검증

    Args:
        token: JWT 토큰

    Returns:
        {"valid": True, "email": "..."} 또는 {"valid": False, "error": "..."}
    """
    if not JWT_SECRET_KEY:
        return {"valid": False, "error": "JWT_SECRET_KEY 환경변수가 설정되지 않았습니다."}

    try:
        payload = jwt.decode(token, JWT_SECRET_KEY, algorithms=[JWT_ALGORITHM])

        # purpose 검증
        if payload.get("purpose") != "email_verification":
            return {"valid": False, "error": "유효하지 않은 토큰입니다."}

        return {"valid": True, "email": payload["email"]}

    except jwt.ExpiredSignatureError:
        return {"valid": False, "error": "인증 토큰이 만료되었습니다. 다시 인증해주세요."}
    except jwt.InvalidTokenError:
        return {"valid": False, "error": "유효하지 않은 인증 토큰입니다."}


@app.post("/api/alerts/send-verification")
@limiter.limit("3/minute")  # 분당 3회 제한 (스팸 방지)
async def send_verification_email(request: Request):
    """
    이메일 인증 요청 - 인증 이메일 발송 (JWT 방식)

    보안: Rate Limit으로 스팸 방지 (분당 3회)
    사용자가 이메일을 입력하고 '인증하기' 버튼을 누르면
    해당 이메일로 JWT 토큰이 포함된 인증 링크를 발송합니다.

    JWT 토큰은 30분간 유효하며, 서버 재시작과 무관하게 검증 가능합니다.
    """
    try:
        body = await request.json()
        email = body.get("email", "").strip()

        # 이메일 형식 검증
        import re

        email_regex = r"^[^\s@]+@[^\s@]+\.[^\s@]+$"
        if not email or not re.match(email_regex, email):
            raise HTTPException(status_code=400, detail="올바른 이메일 주소를 입력해주세요.")

        # 이미 인증된 이메일인지 확인
        state_manager = get_state_manager()
        existing = state_manager.get_subscription(email)
        if existing and existing.verified:
            return {
                "success": True,
                "already_verified": True,
                "message": "이미 인증 완료된 이메일입니다.",
            }

        # JWT 토큰 생성 (30분 유효)
        token = create_email_verification_token(email)

        # 인증 전용 페이지 URL 생성 (대시보드 대신 전용 페이지로 리다이렉트)
        base_url = get_base_url()
        verify_url = f"{base_url}/api/alerts/confirm-email?token={token}&email={email}"

        # EmailSender 직접 사용
        from src.tools.notifications.email_sender import EmailSender

        email_sender = EmailSender()

        if not email_sender.is_enabled():
            raise HTTPException(status_code=503, detail="이메일 서비스가 설정되지 않았습니다.")

        # 인증 이메일 발송
        result = await email_sender.send_verification_email(
            recipient=email, verify_url=verify_url, token=token
        )

        if result.success:
            logging.info(
                f"Verification email sent to {email} (JWT, expires in {EMAIL_VERIFICATION_EXPIRES_MINUTES}min)"
            )
            return {
                "success": True,
                "message": "인증 이메일이 발송되었습니다. (30분 내 인증해주세요)",
            }
        else:
            raise HTTPException(status_code=500, detail=f"이메일 발송 실패: {result.message}")

    except ValueError as e:
        # JWT_SECRET_KEY 미설정 에러
        logging.error(f"JWT configuration error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e
    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Send verification email error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@app.post("/api/alerts/verify-email")
@limiter.limit("10/minute")  # 분당 10회 제한 (brute force 방지)
async def verify_email_token_endpoint(request: Request):
    """
    이메일 인증 토큰 검증 (JWT 방식)

    보안: Rate Limit으로 brute force 방지 (분당 10회)
    사용자가 이메일의 인증 버튼을 클릭하면
    JWT 토큰을 검증하고 이메일 인증 상태를 StateManager에 영구 저장합니다.

    JWT 토큰은 stateless이므로 서버 재시작과 무관하게 검증 가능합니다.
    """
    try:
        body = await request.json()
        token = body.get("token", "")
        email = body.get("email", "").strip()

        if not token or not email:
            raise HTTPException(status_code=400, detail="토큰과 이메일이 필요합니다.")

        # JWT 토큰 검증
        result = verify_jwt_email_token(token)

        if not result["valid"]:
            raise HTTPException(status_code=400, detail=result["error"])

        # 토큰의 이메일과 요청 이메일 일치 확인
        token_email = result["email"]
        if token_email != email:
            raise HTTPException(status_code=400, detail="이메일이 일치하지 않습니다.")

        # StateManager에 인증 완료 상태 영구 저장
        try:
            state_manager = get_state_manager()

            # 기존 구독 정보 확인
            existing = state_manager.get_subscription(email)

            if existing:
                # 기존 구독이 있으면 verified 상태 업데이트 + 활성화
                existing.verified = True
                existing.verified_at = datetime.now()
                existing.consent = True
                existing.consent_date = datetime.now()
                existing.active = True
                state_manager._save_subscriptions()
            else:
                # 새 구독 등록 (verified=True로 생성)
                state_manager.register_email(
                    email=email,
                    consent=True,
                    alert_types=["rank_change", "important_insight", "daily_summary"],
                )
                # verified 상태 추가 설정
                subscription = state_manager.get_subscription(email)
                if subscription:
                    subscription.verified = True
                    subscription.verified_at = datetime.now()
                    state_manager._save_subscriptions()

            logging.info(f"Email verified and saved to StateManager: {email}")
        except Exception as e:
            logging.warning(f"Failed to save verification status: {e}")

        return {"verified": True, "email": email, "message": "이메일 인증이 완료되었습니다!"}

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Verify email error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


@app.get("/api/alerts/confirm-email")
async def confirm_email_page(token: str, email: str):
    """
    이메일 인증 확인 페이지 (GET 요청으로 접근)

    사용자가 이메일의 인증 링크를 클릭하면 이 페이지가 표시됩니다.
    토큰을 검증하고 인증 완료 상태를 저장한 후, 창을 닫아도 되는 안내 페이지를 반환합니다.
    원래 대시보드 탭은 폴링으로 인증 완료를 감지하여 자동으로 다음 단계로 이동합니다.
    """
    from fastapi.responses import HTMLResponse

    # JWT 토큰 검증
    result = verify_jwt_email_token(token)

    if not result["valid"]:
        error_html = f"""
        <!DOCTYPE html>
        <html lang="ko">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>인증 실패 - AMORE Pacific</title>
            <style>
                * {{ margin: 0; padding: 0; box-sizing: border-box; }}
                body {{
                    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                    background: linear-gradient(135deg, #001C58 0%, #1F5795 100%);
                    min-height: 100vh;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    padding: 20px;
                }}
                .card {{
                    background: white;
                    border-radius: 20px;
                    padding: 48px;
                    max-width: 420px;
                    width: 100%;
                    text-align: center;
                    box-shadow: 0 20px 60px rgba(0,0,0,0.3);
                }}
                .icon {{
                    width: 80px;
                    height: 80px;
                    background: #fee2e2;
                    border-radius: 50%;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    margin: 0 auto 24px;
                }}
                .icon svg {{ width: 40px; height: 40px; color: #ef4444; }}
                h1 {{ color: #001C58; font-size: 24px; margin-bottom: 12px; }}
                p {{ color: #64748b; font-size: 15px; line-height: 1.6; }}
                .error-msg {{ color: #ef4444; font-size: 13px; margin-top: 16px; padding: 12px; background: #fef2f2; border-radius: 8px; }}
            </style>
        </head>
        <body>
            <div class="card">
                <div class="icon">
                    <svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                        <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M6 18L18 6M6 6l12 12"/>
                    </svg>
                </div>
                <h1>인증 실패</h1>
                <p>이메일 인증 링크가 만료되었거나 유효하지 않습니다.</p>
                <div class="error-msg">{result.get('error', '토큰이 유효하지 않습니다.')}</div>
                <p style="margin-top: 20px; font-size: 13px;">대시보드에서 다시 인증을 요청해주세요.</p>
            </div>
        </body>
        </html>
        """
        return HTMLResponse(content=error_html, status_code=400)

    # 토큰의 이메일과 요청 이메일 일치 확인
    token_email = result["email"]
    if token_email != email:
        return HTMLResponse(content="이메일이 일치하지 않습니다.", status_code=400)

    # StateManager에 인증 완료 상태 저장
    try:
        state_manager = get_state_manager()
        existing = state_manager.get_subscription(email)

        if existing:
            existing.verified = True
            existing.verified_at = datetime.now()
            existing.consent = True
            existing.consent_date = datetime.now()
            existing.active = True
            state_manager._save_subscriptions()
        else:
            state_manager.register_email(
                email=email,
                consent=True,
                alert_types=["rank_change", "important_insight", "daily_summary"],
            )
            subscription = state_manager.get_subscription(email)
            if subscription:
                subscription.verified = True
                subscription.verified_at = datetime.now()
                state_manager._save_subscriptions()

        logging.info(f"Email verified via confirm page: {email}")
    except Exception as e:
        logging.warning(f"Failed to save verification status: {e}")

    # 인증 성공 페이지 반환
    success_html = f"""
    <!DOCTYPE html>
    <html lang="ko">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>인증 완료 - AMORE Pacific</title>
        <style>
            * {{ margin: 0; padding: 0; box-sizing: border-box; }}
            body {{
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: linear-gradient(135deg, #001C58 0%, #1F5795 100%);
                min-height: 100vh;
                display: flex;
                align-items: center;
                justify-content: center;
                padding: 20px;
            }}
            .card {{
                background: white;
                border-radius: 20px;
                padding: 48px;
                max-width: 420px;
                width: 100%;
                text-align: center;
                box-shadow: 0 20px 60px rgba(0,0,0,0.3);
            }}
            .icon {{
                width: 80px;
                height: 80px;
                background: #d1fae5;
                border-radius: 50%;
                display: flex;
                align-items: center;
                justify-content: center;
                margin: 0 auto 24px;
                animation: pulse 2s infinite;
            }}
            @keyframes pulse {{
                0%, 100% {{ transform: scale(1); }}
                50% {{ transform: scale(1.05); }}
            }}
            .icon svg {{ width: 40px; height: 40px; color: #10b981; }}
            h1 {{ color: #001C58; font-size: 24px; margin-bottom: 12px; }}
            p {{ color: #64748b; font-size: 15px; line-height: 1.6; }}
            .email {{
                color: #1F5795;
                font-weight: 600;
                background: #f0f9ff;
                padding: 8px 16px;
                border-radius: 8px;
                display: inline-block;
                margin: 16px 0;
            }}
            .hint {{
                margin-top: 24px;
                padding: 16px;
                background: #f8fafc;
                border-radius: 12px;
                font-size: 13px;
                color: #475569;
            }}
            .close-btn {{
                margin-top: 24px;
                padding: 14px 32px;
                background: #001C58;
                color: white;
                border: none;
                border-radius: 10px;
                font-size: 15px;
                font-weight: 600;
                cursor: pointer;
                transition: background 0.2s;
            }}
            .close-btn:hover {{ background: #1F5795; }}
        </style>
    </head>
    <body>
        <div class="card">
            <div class="icon">
                <svg xmlns="http://www.w3.org/2000/svg" fill="none" viewBox="0 0 24 24" stroke="currentColor">
                    <path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M5 13l4 4L19 7"/>
                </svg>
            </div>
            <h1>이메일 인증 완료!</h1>
            <div class="email">{email}</div>
            <p>이메일 주소가 성공적으로 인증되었습니다.</p>
            <div class="hint">
                이 창은 닫아도 됩니다.<br>
                원래 대시보드 화면에서 자동으로 다음 단계로 이동합니다.
            </div>
            <button class="close-btn" onclick="window.close()">이 창 닫기</button>
        </div>
    </body>
    </html>
    """
    return HTMLResponse(content=success_html)


@app.get("/api/alerts/verification-status")
async def get_verification_status(email: str):
    """
    이메일 인증 상태 확인 (StateManager 기반)

    JWT 방식으로 변경되어 인증 완료 상태는 StateManager에 영구 저장됩니다.
    """
    try:
        state_manager = get_state_manager()
        subscription = state_manager.get_subscription(email)

        if subscription:
            return {
                "verified": subscription.verified,
                "status": "verified" if subscription.verified else "pending",
                "verified_at": subscription.verified_at.isoformat()
                if subscription.verified_at
                else None,
            }

        return {"verified": False, "status": "not_found"}

    except Exception as e:
        logging.error(f"Get verification status error: {e}")
        return {"verified": False, "status": "error", "error": str(e)}


# ============= Insight Email API =============


@app.post("/api/alerts/send-insight-report")
async def send_insight_report_email(request: Request):
    """
    인사이트 리포트 이메일 발송 (수동)

    대시보드에서 '이메일로 보내기' 버튼 클릭 시 호출됩니다.
    현재 인사이트와 KPI 데이터를 이메일로 발송합니다.

    StateManager 기반 인증 상태 확인 (JWT 방식 변경에 따른 업데이트)
    """
    try:
        body = await request.json()
        recipient_email = body.get("email", "").strip()

        if not recipient_email:
            raise HTTPException(status_code=400, detail="이메일 주소가 필요합니다.")

        # StateManager에서 이메일 인증 상태 확인
        state_manager = get_state_manager()
        subscription = state_manager.get_subscription(recipient_email)

        if not subscription or not subscription.verified:
            raise HTTPException(
                status_code=403, detail="이메일 인증이 필요합니다. 먼저 이메일을 인증해주세요."
            )

        # EmailSender 초기화
        from src.tools.notifications.email_sender import EmailSender

        email_sender = EmailSender()

        if not email_sender.is_enabled():
            raise HTTPException(status_code=503, detail="이메일 서비스가 설정되지 않았습니다.")

        # 현재 대시보드 데이터 로드
        dashboard_data = load_dashboard_data()
        if not dashboard_data:
            raise HTTPException(status_code=404, detail="대시보드 데이터가 없습니다.")

        # KPI 계산
        products = dashboard_data.get("products", [])
        laneige_products = [p for p in products if p.get("brand") == "LANEIGE"]
        avg_rank = (
            sum(p.get("rank", 100) for p in laneige_products) / len(laneige_products)
            if laneige_products
            else 0
        )

        # SoS 계산 (Top 100 기준)
        top100 = products[:100]
        laneige_in_top100 = len([p for p in top100 if p.get("brand") == "LANEIGE"])
        sos = (laneige_in_top100 / len(top100) * 100) if top100 else 0

        # HHI 계산
        brand_counts = {}
        for p in top100:
            brand = p.get("brand", "Unknown")
            brand_counts[brand] = brand_counts.get(brand, 0) + 1
        hhi = (
            sum((count / len(top100) * 100) ** 2 for count in brand_counts.values())
            if top100
            else 0
        )

        # 인사이트 가져오기 (캐시된 것 또는 새로 생성)
        insight_content = dashboard_data.get("latest_insight", "")
        if not insight_content:
            insight_content = (
                "<p>현재 생성된 인사이트가 없습니다. 대시보드에서 인사이트를 먼저 생성해주세요.</p>"
            )
        else:
            # 마크다운을 HTML로 간단 변환
            insight_content = insight_content.replace("\n\n", "</p><p>").replace("\n", "<br>")
            insight_content = f"<p>{insight_content}</p>"

        # Top 10 제품 데이터
        top10_products = []
        for i, p in enumerate(products[:10]):
            top10_products.append(
                {
                    "rank": i + 1,
                    "name": p.get("title", "N/A"),
                    "brand": p.get("brand", "Unknown"),
                    "change": p.get("rank_change", 0),
                }
            )

        # 브랜드별 변동
        brand_changes = []
        for brand in ["LANEIGE", "e.l.f.", "Maybelline", "Summer Fridays", "COSRX"]:
            brand_products = [p for p in products if p.get("brand") == brand]
            if brand_products:
                avg_change = sum(p.get("rank_change", 0) for p in brand_products) / len(
                    brand_products
                )
                if avg_change > 0:
                    brand_changes.append(
                        {
                            "brand": brand,
                            "change_text": f"평균 ▲{avg_change:.1f} 상승",
                            "color": "#28a745",
                        }
                    )
                elif avg_change < 0:
                    brand_changes.append(
                        {
                            "brand": brand,
                            "change_text": f"평균 ▼{abs(avg_change):.1f} 하락",
                            "color": "#dc3545",
                        }
                    )

        # 리포트 날짜
        report_date = datetime.now().strftime("%Y년 %m월 %d일")

        # 대시보드 URL (Railway 자동 감지)
        dashboard_url = get_base_url() + "/dashboard"

        # 이메일 발송
        result = await email_sender.send_insight_report(
            recipients=[recipient_email],
            report_date=report_date,
            avg_rank=avg_rank,
            sos=sos,
            hhi=hhi,
            insight_content=insight_content,
            top10_products=top10_products,
            brand_changes=brand_changes,
            dashboard_url=dashboard_url,
        )

        if result.success:
            logging.info(f"Insight report sent to {recipient_email}")
            return {
                "success": True,
                "message": f"인사이트 리포트가 {recipient_email}로 발송되었습니다.",
                "sent_to": result.sent_to,
            }
        else:
            raise HTTPException(status_code=500, detail=f"이메일 발송 실패: {result.message}")

    except HTTPException:
        raise
    except Exception as e:
        logging.error(f"Send insight report error: {e}")
        raise HTTPException(status_code=500, detail=str(e)) from e


# ============= 서버 실행 =============

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8001)
